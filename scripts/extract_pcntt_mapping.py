from pathlib import Path
import json
import sys

from docx import Document


ROOT_DIR = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT_DIR / "PCNTT_MAPPING_FILE.docx"
OUTPUT_PATH = ROOT_DIR / "storage" / "business_mapping" / "pcntt_mapping.json"

FAQ_TABLE_SOURCES = [
    "2026.03.03.ChatbotAI_CBGV_SV_V4",
    "2026.03.25.AI_HDSD TREN WEB SUPPORT SV",
    "2026.03.25.AI_HDSD TREN WEB SUPPORT CBGV",
]


def _clean(value: str) -> str:
    return " ".join(str(value or "").split())


def _split_keywords(value: str) -> list[str]:
    return [
        keyword.strip()
        for keyword in _clean(value).split(",")
        if keyword.strip()
    ]


def _build_catalog(table) -> dict[str, list[dict]]:
    catalog: dict[str, list[dict]] = {}

    for row in table.rows[1:]:
        cells = [_clean(cell.text) for cell in row.cells]
        if len(cells) < 5 or not cells[0]:
            continue

        item = {
            "file_id": cells[0],
            "source_file": cells[1],
            "document_number": cells[2],
            "source_structure": cells[3],
            "audience": cells[4],
        }
        catalog.setdefault(cells[0], []).append(item)

    return catalog


def _catalog_match(catalog: dict[str, list[dict]], file_id: str, source_file: str) -> dict:
    entries = catalog.get(file_id) or []
    for entry in entries:
        if entry.get("source_file") == source_file:
            return entry

    return entries[0] if entries else {}


def extract_mapping(docx_path: Path = DOCX_PATH) -> dict:
    document = Document(docx_path)
    if len(document.tables) < 2:
        raise ValueError("PCNTT mapping document does not contain FAQ tables")

    catalog = _build_catalog(document.tables[0])
    records = []
    sequence = 1

    for table_index, table in enumerate(document.tables[1:]):
        source_file = FAQ_TABLE_SOURCES[table_index] if table_index < len(FAQ_TABLE_SOURCES) else None

        for row in table.rows[1:]:
            cells = [_clean(cell.text) for cell in row.cells]
            if len(cells) < 6:
                continue

            stt, file_id, question, answer, source_location, keywords = cells[:6]
            if not question or not answer:
                continue

            catalog_item = _catalog_match(catalog, file_id, source_file or "")
            records.append({
                "id": f"PCNTT_MAP_{sequence:04d}",
                "unit": "PCNTT",
                "stt": stt,
                "file_id": file_id,
                "source_file": source_file or catalog_item.get("source_file"),
                "source_location": source_location,
                "question": question,
                "answer": answer,
                "keywords": _split_keywords(keywords),
                "audience": catalog_item.get("audience"),
            })
            sequence += 1

    return {
        "source_docx": docx_path.name,
        "record_count": len(records),
        "records": records,
    }


def main():
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    payload = extract_mapping()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    with OUTPUT_PATH.open("w", encoding="utf-8") as file:
        json.dump(payload, file, ensure_ascii=False, indent=2)
        file.write("\n")

    print(f"Wrote {payload['record_count']} records to {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
