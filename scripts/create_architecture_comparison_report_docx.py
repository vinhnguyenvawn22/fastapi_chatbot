from pathlib import Path
import argparse
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT_DIR = Path(__file__).resolve().parent.parent
SOURCE_FILE = ROOT_DIR / "BAO_CAO_KIEN_TRUC_VA_SO_SANH_HAI_CHATBOT.md"
OUTPUT_FILE = ROOT_DIR / "BAO_CAO_KIEN_TRUC_VA_SO_SANH_HAI_CHATBOT.docx"

NAVY = "17365D"
BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"
TEXT = RGBColor(31, 31, 31)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_heading(paragraph, keep_with_next=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_with_next:
        keep_next = OxmlElement("w:keepNext")
        p_pr.append(keep_next)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Trang ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_run_font(run, size=12, bold=None, italic=None, color=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_inline_runs(paragraph, text: str, size=12, color=TEXT) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(
                run,
                size=max(size - 1, 8),
                color=RGBColor(128, 0, 0),
                name="Consolas",
            )
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, color=color)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, color in (
        ("Title", 22, NAVY),
        ("Heading 1", 16, NAVY),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 12, BLUE),
    ):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)


def add_cover(
    doc: Document,
    title_text: str = "BÁO CÁO KIẾN TRÚC CHATBOT RAG\nVÀ SO SÁNH HIỆU QUẢ HAI CHATBOT",
    subtitle_text: str = "Trọng tâm: Endpoint POST /api/chat/local-documents",
    intro_first: bool = False,
) -> None:
    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    set_run_font(run, size=22, bold=True, color=RGBColor.from_string(NAVY))
    title.paragraph_format.space_after = Pt(24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(subtitle_text)
    set_run_font(run, size=14, bold=True, color=RGBColor.from_string(BLUE))

    doc.add_paragraph()
    meta = (
        [
            ("Phạm vi kỹ thuật", "Endpoint POST /api/chat/local-documents"),
            ("Phương pháp", "Retrieval-Augmented Generation (RAG)"),
            ("Nội dung", "Giới thiệu kiến trúc, luồng xử lý và đánh giá thực nghiệm"),
            ("Ngày lập báo cáo", "30/07/2026"),
        ]
        if intro_first
        else [
            ("Đối tượng so sánh", "Chatbot 1 của Trường và Chatbot 2 của đề tài"),
            ("Dữ liệu đánh giá", "60 câu hỏi Ground Truth, 120 lượt chấm"),
            ("Thang điểm", "Đúng nguồn, đúng trọng tâm, đủ ý, dễ hiểu, không bịa"),
            ("Ngày lập báo cáo", "30/07/2026"),
        ]
    )
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, (label, value) in enumerate(meta):
        left, right = table.rows[index].cells
        left.width = Cm(4.5)
        right.width = Cm(10.5)
        set_cell_shading(left, LIGHT_BLUE)
        set_cell_margins(left)
        set_cell_margins(right)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(label)
        set_run_font(r, size=11, bold=True)
        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(value)
        set_run_font(r, size=11)

    doc.add_paragraph()
    conclusion = doc.add_paragraph()
    conclusion.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conclusion_text = (
        "Phần đầu báo cáo trình bày độc lập mục tiêu, kiến trúc và luồng xử lý "
        "của hệ thống trước khi chuyển sang phần đánh giá thực nghiệm."
        if intro_first
        else (
            "Kết luận chính: Chatbot 2 đạt 504,5/600 điểm, "
            "cao hơn Chatbot 1 89,5 điểm."
        )
    )
    run = conclusion.add_run(conclusion_text)
    set_run_font(run, size=13, bold=True, color=RGBColor.from_string(NAVY))
    doc.add_page_break()


def add_contents(doc: Document, markdown: str) -> None:
    heading = doc.add_paragraph("NỘI DUNG BÁO CÁO", style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    items = []
    for line in markdown.splitlines():
        match = re.match(r"^(#{1,2})\s+(.+)$", line.strip())
        if not match:
            continue
        text = match.group(2).strip()
        if text.startswith("BÁO CÁO"):
            continue
        if text.startswith("PHẦN ") or re.match(r"^\d+\.\d+\.", text):
            items.append((len(match.group(1)), text))
    for level, item in items:
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Cm(0.2 if level == 1 else 0.8)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(2)
        add_inline_runs(p, item, size=12)
        if level == 1:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_page_break()


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    index = start
    while index < len(lines):
        line = lines[index].strip()
        if not (line.startswith("|") and line.endswith("|")):
            break
        if not is_table_separator(line):
            rows.append([cell.strip() for cell in line.strip("|").split("|")])
        index += 1
    return rows, index


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for row_index, values in enumerate(rows):
        for col_index in range(columns):
            cell = table.cell(row_index, col_index)
            set_cell_margins(cell, top=65, start=70, bottom=65, end=70)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, NAVY)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            text = values[col_index] if col_index < len(values) else ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if col_index == 0
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            add_inline_runs(
                paragraph,
                text,
                size=9,
                color=RGBColor(255, 255, 255) if row_index == 0 else TEXT,
            )
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code_block(doc: Document, code: str, language: str = "") -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.4)
    paragraph.paragraph_format.right_indent = Cm(0.4)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EAF0F6")
    p_pr.append(shd)
    label = "Sơ đồ luồng (dạng mã Mermaid)\n" if language == "mermaid" else ""
    run = paragraph.add_run(label + code.strip())
    set_run_font(run, size=8.5, name="Consolas", color=RGBColor(35, 35, 35))


def add_blockquote(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.8)
    paragraph.paragraph_format.right_indent = Cm(0.5)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    add_inline_runs(paragraph, text, size=11, color=RGBColor.from_string(BLUE))
    for run in paragraph.runs:
        run.italic = True


def add_markdown_body(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    # Skip the Markdown title/metadata block; the Word file has its own cover.
    first_rule = next(
        (index for index, line in enumerate(lines) if line.strip() == "---"),
        -1,
    )
    index = first_rule + 1 if first_rule >= 0 else 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer).strip()
        paragraph_buffer.clear()
        if text:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline_runs(paragraph, text, size=12)

    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            language = stripped[3:].strip()
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(doc, "\n".join(code_lines), language)
            index += 1
            continue

        if (
            stripped.startswith("|")
            and stripped.endswith("|")
            and index + 1 < len(lines)
            and is_table_separator(lines[index + 1].strip())
        ):
            flush_paragraph()
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            paragraph = doc.add_paragraph(
                style={1: "Heading 1", 2: "Heading 1", 3: "Heading 2"}[level]
            )
            add_inline_runs(
                paragraph,
                heading_text,
                size={1: 16, 2: 16, 3: 14}[level],
                color=RGBColor.from_string(NAVY if level <= 2 else BLUE),
            )
            for run in paragraph.runs:
                run.bold = True
            set_repeat_heading(paragraph)
            index += 1
            continue

        if stripped == "---":
            flush_paragraph()
            index += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            add_blockquote(doc, stripped.lstrip(">").strip())
            index += 1
            continue

        bullet_match = re.match(r"^\s*-\s+(.+)$", line)
        if bullet_match:
            flush_paragraph()
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, bullet_match.group(1), size=12)
            index += 1
            continue

        number_match = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if number_match:
            flush_paragraph()
            paragraph = doc.add_paragraph(style="List Number")
            add_inline_runs(paragraph, number_match.group(1), size=12)
            index += 1
            continue

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()


def add_header_footer(
    doc: Document,
    header_text: str = "Báo cáo kiến trúc RAG và so sánh hai chatbot",
) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run(header_text)
        set_run_font(run, size=9, italic=True, color=RGBColor.from_string(BLUE))
        add_page_number(section.footer.paragraphs[0])


def set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Báo cáo kiến trúc Chatbot RAG và so sánh hiệu quả hai chatbot"
    props.subject = "Endpoint POST /api/chat/local-documents"
    props.author = "Nhóm thực hiện đề tài"
    props.keywords = "RAG, FastAPI, chatbot, hybrid retrieval, BM25, ChromaDB, Gemini"


def build_report(
    source_file: Path = SOURCE_FILE,
    output_file: Path = OUTPUT_FILE,
    title_text: str = "BÁO CÁO KIẾN TRÚC CHATBOT RAG\nVÀ SO SÁNH HIỆU QUẢ HAI CHATBOT",
    intro_first: bool = False,
    include_contents: bool = True,
) -> Path:
    markdown = source_file.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_cover(doc, title_text=title_text, intro_first=intro_first)
    if include_contents:
        add_contents(doc, markdown)
    add_markdown_body(doc, markdown)
    add_header_footer(
        doc,
        header_text=(
            "Báo cáo giới thiệu và đánh giá chatbot"
            if intro_first
            else "Báo cáo kiến trúc RAG và so sánh hai chatbot"
        ),
    )
    set_core_properties(doc)
    doc.save(output_file)
    return output_file


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=SOURCE_FILE)
    parser.add_argument("--output", type=Path, default=OUTPUT_FILE)
    parser.add_argument(
        "--title",
        default="BÁO CÁO KIẾN TRÚC CHATBOT RAG\nVÀ SO SÁNH HIỆU QUẢ HAI CHATBOT",
    )
    parser.add_argument("--intro-first", action="store_true")
    parser.add_argument("--no-contents", action="store_true")
    args = parser.parse_args()
    output = build_report(
        source_file=args.source.resolve(),
        output_file=args.output.resolve(),
        title_text=args.title,
        intro_first=args.intro_first,
        include_contents=not args.no_contents,
    )
    print(output)
