from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


OUTPUT = Path("THIET_KE_ENDPOINT_RAG_THONG_NHAT_BM25_ANN.docx")
FONT = "Times New Roman"


def apply_font(run, size=11, bold=None, italic=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_title(doc):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("THIẾT KẾ ENDPOINT RAG THỐNG NHẤT BM25 + ANN")
    apply_font(run, size=16, bold=True, color=(31, 78, 121))

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Dự án FastAPI Chatbot RAG UNETI")
    apply_font(run, size=12, italic=True)


def heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        apply_font(run, color=(31, 78, 121))


def paragraph(doc, text=""):
    p = doc.add_paragraph()
    run = p.add_run(text)
    apply_font(run)
    p.paragraph_format.space_after = Pt(4)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        apply_font(run)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        apply_font(run)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        cell = t.rows[0].cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, "D9EAF7")
        run = cell.paragraphs[0].add_run(header)
        apply_font(run, bold=True)

    for row in rows:
        cells = t.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            run = cell.paragraphs[0].add_run(str(value))
            apply_font(run)
    doc.add_paragraph()


def build_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(11)

    add_title(doc)
    paragraph(
        doc,
        "Mục tiêu tài liệu: mô tả các kỹ thuật và luồng logic đang dùng trong chatbot hiện tại, "
        "đồng thời đặc tả một endpoint mới gom toàn bộ tài liệu vào một kho tìm kiếm thống nhất. "
        "Endpoint mới không tìm kiếm song song nội bộ/nghiệp vụ và không dùng file mapping "
        "PCNTT_MAPPING_FILE.docx; ý tưởng chủ đạo vẫn là truy xuất bằng BM25 và ANN, sau đó "
        "hợp nhất/rerank/chọn context để sinh câu trả lời.",
    )

    heading(doc, "1. Tóm tắt ý tưởng endpoint mới")
    paragraph(doc, "Endpoint mới đề xuất: POST /api/chat/unified hoặc POST /api/chat/all-documents.")
    paragraph(
        doc,
        "Nguyên tắc chính: thay vì phân route rồi gọi retrieve_business và retrieve_internal song song, "
        "endpoint mới coi tất cả tài liệu là một corpus thống nhất. Câu hỏi chỉ đi qua một pipeline "
        "retrieval duy nhất: chuẩn hóa câu hỏi -> BM25 -> ANN/vector -> fusion -> rerank -> chọn context "
        "-> gọi Gemini để trả lời.",
    )
    bullets(doc, [
        "Không đọc, không chấm điểm, không ưu tiên PCNTT_MAPPING_FILE.docx.",
        "Không tạo nhánh tìm kiếm song song business/internal trong controller.",
        "Tài liệu nội bộ, tài liệu nghiệp vụ, website đã index nếu muốn dùng đều được đưa vào cùng một index với metadata source_type.",
        "BM25 bắt keyword, chính tả, mã văn bản, điều/mục và thuật ngữ chính xác.",
        "ANN bắt ngữ nghĩa tương đồng khi người dùng hỏi khác cách diễn đạt trong tài liệu.",
        "Fusion/rerank quyết định chunk nào vào prompt, không route thắng sớm theo nhóm tài liệu.",
    ])

    heading(doc, "2. Endpoint hiện tại trong dự án")
    table(doc, ["Endpoint", "Handler", "Luồng chính", "Ghi chú"], [
        ["/api/chat/", "handle_chat()", "Tự phân loại, có thể tổng hợp business + internal + website.", "Nhiều rule chọn route, multi-hop và fallback."],
        ["/api/chat/internal", "handle_internal_chat()", "Ép dùng tài liệu nội bộ official_document qua retrieve_internal().", "Cấu trúc gọn, phù hợp làm mẫu cho endpoint mới."],
        ["/api/chat/business", "handle_business_chat()", "Ép dùng tài liệu nghiệp vụ qua retrieve_business().", "Hiện ưu tiên PCNTT_MAPPING_FILE.docx trước generic search."],
        ["/api/chat/website", "handle_website_chat()", "Ép dùng website UNETI.", "Có bước index website rồi search source_type=website_uneti."],
    ])

    heading(doc, "3. Luồng endpoint nội bộ cũ")
    paragraph(doc, "Endpoint /api/chat/internal là mẫu nên tái sử dụng vì ít phân nhánh và không phụ thuộc mapping nghiệp vụ.")
    numbered(doc, [
        "chat_router.py nhận ChatRequest và gọi ConversationService.chat().",
        "ConversationService xử lý session, thread_id, request_id, idempotency và lịch sử hội thoại.",
        "handle_internal_chat() tạo trace, kiểm tra câu hỏi rỗng, phân tích ambiguity/retrieval decision.",
        "_answer_with_internal_documents() gọi retrieve_internal().",
        "retrieve_internal() gọi search_documents(question, source_type_filter=\"official_document\").",
        "search_documents() chạy metadata search, BM25, ANN, RRF, cross-encoder rerank.",
        "Controller kiểm tra confident evidence bằng _has_confident_evidence().",
        "Nếu đủ evidence, generate_answer() build context/prompt rồi gọi Gemini.",
        "Finalize response gồm answer, source, sources, intent, trace_id.",
    ])

    heading(doc, "4. Luồng nghiệp vụ hiện tại và vấn đề mapping")
    paragraph(
        doc,
        "Luồng nghiệp vụ hiện tại trong business_knowledge.py có một tầng PCNTT_MAPPING_FILE.docx đóng vai trò "
        "FAQ/catalog để định hướng source file, location, keyword và audience. Đây là lý do nhiều câu nghiệp vụ "
        "trước tiên được so với mapping, sau đó mới tìm trong tài liệu thật.",
    )
    numbered(doc, [
        "_load_business_index() đọc documents/nghiep_vu; nếu gặp PCNTT_MAPPING_FILE.docx thì build các row source_type=business_faq_mapping.",
        "search_business_sources() tạo mapping_candidates() từ các FAQ row.",
        "Mapping candidate được kiểm tra bằng rule/keyword/audience, có thể dùng mapping judge LLM nếu bật.",
        "Nếu mapping được chọn, hệ thống tìm trong source file theo location, keyword hoặc vector.",
        "Nếu không có mapping tốt, mới fallback sang generic hybrid: keyword/BM25-like + vector runtime + merge score.",
    ])
    paragraph(
        doc,
        "Vấn đề: mapping có thể sai location, chồng rule theo audience/source, và làm endpoint phụ thuộc chất lượng "
        "PCNTT_MAPPING_FILE.docx. Vì vậy endpoint mới nên bỏ mapping layer và search trực tiếp trên tài liệu thật.",
    )

    heading(doc, "5. Kỹ thuật retrieval đang dùng")
    table(doc, ["Kỹ thuật", "File/Module", "Vai trò"], [
        ["Chuẩn hóa text", "query_analyzer.py, elasticsearch_client.py, business_knowledge.py", "Bỏ dấu, lowercase, tách keyword, giảm nhiễu do cách viết khác nhau."],
        ["Query expansion rule", "elasticsearch_client.py", "Mở rộng các cụm như GPA -> điểm trung bình tích lũy, hủy học phần -> rút bớt học phần."],
        ["Metadata search", "elasticsearch_client.py", "Ưu tiên số văn bản, điều, mục, chương, ngày, loại văn bản khi câu hỏi có ràng buộc rõ."],
        ["BM25", "elasticsearch_client.py", "Tìm kiếm lexical trên corpus local, tốt cho từ khóa chính xác, mã văn bản, thuật ngữ hiếm."],
        ["ANN/vector search", "vector_store.py + embedding_client.py", "Tìm chunk gần nghĩa bằng embedding trong ChromaDB."],
        ["RRF Fusion", "elasticsearch_client.py", "Gộp nhiều danh sách kết quả theo thứ hạng thay vì phụ thuộc một score duy nhất."],
        ["Cross-encoder rerank", "reranker.py", "Chấm lại top candidates theo cặp question-document để tăng độ đúng của top source."],
        ["Evidence check", "chatbot_controller.py", "Chặn trả lời khi nguồn không đủ liên quan."],
        ["Context builder", "prompt_builder.py", "Đóng gói các chunk thành thẻ <NGUON> có metadata để Gemini tổng hợp."],
        ["Gemini fallback", "langchain_pipeline.py", "Nếu Gemini lỗi quota/rate limit, dùng fallback extractive từ các câu trong docs."],
        ["Trace logging", "trace_logger.py", "Ghi lại từng bước retrieval/generation để debug."],
    ])

    heading(doc, "6. Thiết kế endpoint RAG thống nhất")
    paragraph(
        doc,
        "Endpoint mới bám cấu trúc /api/chat/internal nhưng thay retrieve_internal() bằng retrieve_unified_documents(). "
        "Hàm này gọi search_documents() không truyền source_type_filter, tức source_type_filter=None, để tất cả source_type "
        "cùng được xét trong một pipeline.",
    )
    numbered(doc, [
        "Router thêm POST /api/chat/unified dùng ChatRequest/ChatResponse giống các endpoint chat khác.",
        "Controller thêm handle_unified_chat() tương tự handle_internal_chat(): nhận question, tạo trace, phân tích ambiguity, gọi _answer_with_unified_documents().",
        "_answer_with_unified_documents() giống _answer_with_internal_documents(), nhưng reason là explicit_unified_endpoint.",
        "Pipeline thêm retrieve_unified() gọi search_documents(question, source_type_filter=None).",
        "search_documents() dùng chung index document: metadata search + BM25 + ANN + RRF + cross-encoder rerank.",
        "Sau khi có docs, controller vẫn dùng _has_confident_evidence(), _build_sources(), generate_answer() và _finalize() như endpoint nội bộ cũ.",
    ])

    heading(doc, "7. Điều cần thay đổi ở dữ liệu/index")
    paragraph(
        doc,
        "Để endpoint thống nhất hoạt động đúng, mọi tài liệu cần được đưa vào cùng cơ chế index. Không nên để tài liệu "
        "nghiệp vụ chỉ nằm trong business_knowledge.py runtime index riêng nếu endpoint mới gọi search_documents().",
    )
    bullets(doc, [
        "Tài liệu nội bộ/quy chế: source_type=official_document.",
        "Tài liệu nghiệp vụ/Web Support: source_type=business_document.",
        "Website UNETI nếu dùng chung: source_type=website_uneti.",
        "Mỗi chunk cần có doc_name, title, content, relative_path, source_type, chunk_index, source_root và metadata bổ sung nếu có.",
        "PCNTT_MAPPING_FILE.docx nên bị loại khỏi endpoint mới, hoặc index với source_type riêng và exclude khỏi retrieval.",
        "Nên reindex ChromaDB để tất cả chunk cần tìm có embedding sẵn; tránh runtime embedding nhiều như generic business search hiện tại.",
    ])

    heading(doc, "8. BM25 + ANN trong endpoint mới")
    numbered(doc, [
        "Normalize/rewrite nhẹ câu hỏi bằng rule, không gọi Gemini nếu không cần.",
        "BM25 search trên toàn corpus để lấy top BM25 candidates.",
        "ANN search trên ChromaDB cùng corpus để lấy top vector candidates.",
        "Gộp kết quả bằng RRF hoặc weighted fusion.",
        "Áp dụng metadata/source scoring mềm: ưu tiên official_document cho quy định, business_document cho thao tác hệ thống, website_uneti cho tin tức/thông báo/link.",
        "Cross-encoder rerank top N candidates.",
        "Chọn context đa dạng theo doc_name/source_type/chủ đề để tránh chỉ lấy nhiều chunk từ một file.",
        "Build prompt và sinh câu trả lời cuối.",
    ])

    heading(doc, "9. Những phần phải bỏ trong endpoint mới")
    bullets(doc, [
        "Không gọi _mapping_candidates().",
        "Không gọi _score_business_faq().",
        "Không dùng BUSINESS_FAQ_SOURCE_TYPE làm nguồn ứng viên.",
        "Không gọi _mapping_gate_decision() hoặc _judge_mapping_with_llm().",
        "Không dùng faq_location để tìm _location_windows().",
        "Không dùng các override sửa sai mapping như survey_source_override hay procedure_evaluation_location_override cho endpoint mới.",
        "Không chạy retrieve_business và retrieve_internal song song trong controller.",
    ])

    heading(doc, "10. Pseudo-code endpoint mới")
    paragraph(doc, "Router: POST /api/chat/unified -> ConversationService.chat(..., handle_unified_chat).")
    paragraph(doc, "Controller: handle_unified_chat(request) giống handle_internal_chat(), nhưng forced_route=\"unified_document\" và gọi _answer_with_unified_documents().")
    paragraph(doc, "Pipeline: retrieve_unified(state) gọi search_documents(state[\"question\"], debug=debug, source_type_filter=None, ambiguity_decision=state.get(\"ambiguity_decision\")).")
    paragraph(doc, "Generation: dùng generate_answer() hiện có, context lấy từ docs đã rerank.")

    heading(doc, "11. Ưu điểm so với luồng song song/mapping")
    bullets(doc, [
        "Đơn giản hơn: một endpoint, một retrieval pipeline, một cơ chế scoring.",
        "Giảm phụ thuộc file mapping và giảm lỗi do mapping sai location.",
        "Có thể trả lời câu hỏi giao thoa quy định + thao tác hệ thống mà không cần route thắng sớm.",
        "Dễ benchmark vì tất cả nguồn được chấm trong cùng không gian BM25/ANN/rerank.",
        "Dễ debug trace: top BM25, top ANN, RRF, rerank đều trên cùng corpus.",
    ])

    heading(doc, "12. Rủi ro và cách kiểm soát")
    table(doc, ["Rủi ro", "Nguyên nhân", "Cách kiểm soát"], [
        ["Nguồn website lấn nguồn quy chế", "Website có từ khóa giống câu hỏi nhưng không phải căn cứ chính thức.", "Thêm source_type prior: câu hỏi quy định/chế tài ưu tiên official_document."],
        ["Nguồn nghiệp vụ lấn nguồn nội bộ", "Web Support mô tả thao tác, không phải quy định.", "Prompt và rerank metadata phân biệt thao tác hệ thống với căn cứ pháp quy."],
        ["Câu ngắn như GPA tìm yếu", "Thiếu keyword mở rộng.", "Giữ dictionary rewrite nhẹ trước retrieval."],
        ["Top chunk cùng một file quá nhiều", "BM25/ANN cùng kéo một vùng tài liệu.", "Áp dụng diverse selection theo doc_name/source_type/chunk adjacency."],
        ["ANN chậm nếu chưa index sẵn", "Runtime embedding nhiều chunk.", "Reindex trước toàn bộ corpus vào ChromaDB."],
        ["Không còn mapping curated", "Mất lợi thế FAQ chỉ đúng vị trí.", "Bù bằng metadata tốt, keyword expansion và rerank."],
    ])

    heading(doc, "13. Tiêu chí kiểm thử endpoint mới")
    bullets(doc, [
        "Câu quy định đại học: cảnh báo học tập, F/F+, chuyển trường, tốt nghiệp phải ưu tiên QĐ 832 nếu có căn cứ.",
        "Câu thao tác Web Support: khảo sát, đánh giá thủ tục, phúc khảo, lịch học phải lấy tài liệu hướng dẫn thao tác phù hợp.",
        "Câu kết hợp: nghỉ học có phép/không phép, hoãn thi, thi lại phải có thể lấy nhiều nguồn nếu cần.",
        "Trace phải hiển thị top BM25, top ANN, RRF, rerank và final_sources.",
        "Không có final source nào là PCNTT_MAPPING_FILE.docx.",
        "Số lần gọi Gemini mục tiêu: chỉ gọi ở final answer; rewrite/HyDE/retrieval plan nên tắt hoặc cache mạnh cho endpoint này.",
    ])

    heading(doc, "14. Kết luận")
    paragraph(
        doc,
        "Endpoint mới nên được xem là một thử nghiệm RAG thống nhất: bỏ mapping, bỏ tìm kiếm song song, đưa tất cả "
        "tài liệu thật vào một corpus, sau đó dùng BM25 + ANN + fusion + rerank để chọn căn cứ. Thiết kế này phù hợp "
        "nếu mục tiêu là giảm rule chồng chéo và tiến gần hơn tới kiểu search rộng trên một kho tri thức chung rồi "
        "tổng hợp từ nhiều nguồn.",
    )

    for section in doc.sections:
        section.top_margin = Pt(56)
        section.bottom_margin = Pt(56)
        section.left_margin = Pt(56)
        section.right_margin = Pt(56)

    return doc


if __name__ == "__main__":
    build_document().save(OUTPUT)
    print(OUTPUT.resolve())
