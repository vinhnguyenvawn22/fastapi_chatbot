from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT = Path("BAO_CAO_CAI_TIEN_CHATBOT_TU_API_LUAN_PHIEN.docx")


def set_run_font(run, size=12, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")


def set_paragraph_runs(paragraph, size=12):
    for run in paragraph.runs:
        set_run_font(run, size=size)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    set_run_font(run, size=14 if level == 1 else 13, bold=True)
    return paragraph


def add_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(item)
        set_run_font(run)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    set_run_font(run, size=11, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            set_cell_text(cell, value, bold=row_index == 0)
            if row_index == 0:
                shade_cell(cell, "D9EAF7")
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.8)

    for style_name in ("Normal", "List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style.element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")


def build_report():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BÁO CÁO CẢI TIẾN CHATBOT RAG UNETI")
    set_run_font(run, size=16, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Giai đoạn từ commit thêm API key luân phiên đến hiện tại")
    set_run_font(run, size=13, italic=True)

    add_heading(doc, "1. Phạm vi báo cáo")
    add_paragraph(
        doc,
        "Báo cáo này tổng hợp các cải tiến của chatbot RAG UNETI từ commit bắt đầu "
        "triển khai cơ chế sử dụng nhiều Gemini API key luân phiên đến trạng thái hiện tại "
        "của nhánh develop.",
    )
    add_bullets(
        doc,
        [
            "Commit bắt đầu: 9d8df59050ff9124ee46f267b910cf295ccb623d - update chatbot.",
            "Commit cha dùng để so sánh: a804e31b14c369c5a1e5e54e9ef6b0655619ab12.",
            "Phạm vi thay đổi: 21 file, khoảng 4640 dòng thêm và 1153 dòng xóa.",
            "Các nhóm file chính: cấu hình, Gemini client, controller API chat, retrieval nội bộ, business retrieval, prompt/context và test.",
        ],
    )

    add_heading(doc, "2. Bối cảnh trước khi cải tiến")
    add_bullets(
        doc,
        [
            "Chỉ cấu hình một Gemini API key chính, dễ bị gián đoạn khi key hết quota hoặc bị rate limit.",
            "Một câu hỏi có thể gọi Gemini nhiều lần ở các bước phụ như query rewriting, retrieval plan, HyDE và final answer.",
            "Một số câu hỏi học vụ bị chọn sai nguồn, ví dụ lấy tài liệu thạc sĩ hoặc Web Support cho câu hỏi sinh viên đại học.",
            "Các câu hỏi đời thường như GPA, chuyển trường, hủy học phần, F/F+ chưa luôn khớp tốt với thuật ngữ trong quy chế.",
            "Câu trả lời đôi khi chỉ dựa trên một nguồn hoặc bỏ qua chunk có căn cứ rõ.",
            "Chưa có đủ test regression cho các lỗi thực tế phát hiện trong quá trình chấm và test chatbot.",
        ],
    )

    add_heading(doc, "3. Cải tiến sử dụng nhiều Gemini API key luân phiên")
    add_paragraph(
        doc,
        "Đây là nhóm cải tiến mở đầu của giai đoạn báo cáo, nằm trong commit 9d8df59. "
        "Các file liên quan gồm app/core/config.py, app/data/gemini_client.py, .env.example và tests/test_gemini_client.py.",
    )
    add_heading(doc, "3.1. Hỗ trợ nhiều cách khai báo API key", 2)
    add_bullets(
        doc,
        [
            "GEMINI_API_KEYS=key_1,key_2,key_3.",
            "GEMINI_API_KEY_1=key_1, GEMINI_API_KEY_2=key_2, ...",
            "Nhiều dòng GEMINI_API_KEY lặp lại trong file .env thật.",
            "Tự loại bỏ key rỗng, key placeholder như your_api_key_here, changeme, replace_me.",
            "Tự khử trùng lặp để tránh gọi lặp cùng một key.",
        ],
    )
    add_heading(doc, "3.2. Cơ chế round-robin và fallback", 2)
    add_paragraph(
        doc,
        "Gemini client tạo danh sách key theo thứ tự xoay vòng. Mỗi lần gọi Gemini, key bắt đầu "
        "sẽ dịch sang key kế tiếp. Nếu key hiện tại gặp lỗi quota/rate limit hoặc dịch vụ tạm thời "
        "không sẵn sàng, hệ thống thử key tiếp theo.",
    )
    add_bullets(
        doc,
        [
            "Lỗi được fallback: 429, RESOURCE_EXHAUSTED, rate limit, quota, 503, UNAVAILABLE.",
            "Client Gemini được cache theo từng key để không khởi tạo lại liên tục.",
            "Có Lock để việc tạo client an toàn hơn khi nhiều request cùng chạy.",
            "Bộ đếm số lần gọi Gemini vẫn được duy trì theo từng request để debug trace.",
        ],
    )

    add_heading(doc, "4. Giảm số lần gọi Gemini")
    add_bullets(
        doc,
        [
            "Multi-hop retrieval vẫn giữ, nhưng sub-question không gọi Gemini retrieval plan.",
            "Khi query_context có skip_retrieval_plan_llm=True, business retrieval dùng rule/keyword/vector thay vì gọi LLM tạo plan.",
            "Mapping judge LLM được bỏ qua trong sub-question để tránh nhân số lần gọi theo số sub-question.",
            "Một số câu có căn cứ rõ dùng fallback hoặc deterministic answer, giảm phụ thuộc final generation.",
            "Trace ghi gemini_call_count để kiểm tra một request đã gọi Gemini bao nhiêu lần.",
        ],
    )

    add_heading(doc, "5. Cải tiến RAG tổng hợp trên API /api/chat/")
    add_bullets(
        doc,
        [
            "Thêm cơ chế query decomposition rule-first cho câu hỏi so sánh/tổng hợp.",
            "Retrieve song song internal và business khi câu hỏi có thể cần cả quy định và thao tác hệ thống.",
            "Dedupe source theo doc_name, title, chunk_index, source_type.",
            "Chọn nguồn dựa trên độ phủ ý, loại nguồn và độ phù hợp nghiệp vụ thay vì chỉ lấy top score chung.",
            "Giới hạn dominance của một tài liệu để câu trả lời có thể dùng nhiều nguồn khi cần.",
        ],
    )

    add_heading(doc, "6. Cải tiến routing nội bộ, nghiệp vụ và website")
    add_bullets(
        doc,
        [
            "Câu hỏi quy chế/quy định/học vụ ưu tiên source_type=official_document.",
            "Câu hỏi thao tác hệ thống ưu tiên business_document hoặc business_faq_mapping.",
            "Website chỉ là fallback khi không có căn cứ đủ tốt từ nội bộ hoặc nghiệp vụ.",
            "Internal retrieval truyền filter official_document để tránh lấy nhầm website_uneti trong ChromaDB.",
            "Business retrieval vẫn giữ mapping/guided search cho các câu hỏi thao tác.",
        ],
    )

    add_heading(doc, "7. Cải tiến query expansion và policy profile học vụ")
    add_heading(doc, "7.1. Query expansion", 2)
    add_bullets(
        doc,
        [
            "GPA -> điểm trung bình tích lũy, điểm trung bình học tập, điểm hệ 4.",
            "Hủy học phần -> hủy đăng ký học phần, rút bớt học phần, Điều 10, Điều 9.",
            "Chuyển trường -> điều kiện chuyển trường, Hiệu trưởng, cùng ngành, nơi cư trú, Điều 28.",
            "Cảnh báo học tập -> đăng ký khối lượng học tập, không quá 16 tín chỉ, Điều 9.",
            "F/F+ -> thang điểm, điểm chữ, học lại, học đổi, Điều 16 và Điều 11.",
            "Tín chỉ -> 15 tiết lý thuyết, 30 tiết thực hành, 45-60 giờ, Điều 2.",
        ],
    )
    add_heading(doc, "7.2. Policy profile", 2)
    add_bullets(
        doc,
        [
            "grade_average: câu hỏi về GPA và điểm trung bình.",
            "course_registration_change: hủy/rút học phần đã đăng ký.",
            "credit_load_warning: sinh viên bị cảnh báo học tập được đăng ký tối đa bao nhiêu tín chỉ.",
            "transfer_school: chuyển trường, phân biệt với chuyển chương trình đào tạo.",
            "elective_failed_course: học phần tự chọn bị F/F+.",
            "f_grade_comparison: so sánh F+ và F.",
            "credit_definition: một tín chỉ tương đương bao nhiêu tiết/giờ.",
            "attendance_exam_eligibility và absence_permission_comparison: nghỉ học, điểm chuyên cần, cấm thi, nghỉ có phép/không phép.",
        ],
    )

    add_heading(doc, "8. Cải tiến scoring và chọn nguồn")
    add_bullets(
        doc,
        [
            "Cộng điểm cho nguồn đúng Điều/Mục, ví dụ Điều 9, Điều 10, Điều 11, Điều 16, Điều 28.",
            "Cộng điểm mạnh cho Quy chế đào tạo đại học chính quy 832 khi câu hỏi là sinh viên đại học.",
            "Trừ điểm tài liệu thạc sĩ nếu câu hỏi không nhắc thạc sĩ.",
            "Trừ điểm nguồn Web Support, thiết bị, phòng học, nghiên cứu khoa học khi câu hỏi thuộc quy chế học vụ.",
            "Chặn false positive kiểu câu hỏi cảnh báo học tập nhưng nguồn lại nói thời khóa biểu/lịch học.",
            "Không thêm nguồn yếu chỉ để đủ số lượng.",
        ],
    )

    add_heading(doc, "9. Một số lỗi thực tế đã sửa")
    add_table(
        doc,
        [
            ["Câu hỏi", "Vấn đề trước đó", "Cách cải thiện"],
            [
                "Em đang bị cảnh báo học tập thì tối đa được đăng ký bao nhiêu tín chỉ?",
                "Bot trả theo mức chung 3/2 số tín chỉ, bỏ qua khoản riêng 16 tín chỉ.",
                "Thêm profile credit_load_warning và deterministic answer khi có căn cứ “không quá 16 tín chỉ”.",
            ],
            [
                "Tôi muốn chuyển trường, không phải chuyển chương trình đào tạo",
                "Tìm được Điều 28 nhưng bị rule lọc rơi rồi fallback sang website.",
                "Sửa rule lọc “tuyển sinh” chỉ áp dụng trên metadata/tên tài liệu, không loại chunk đúng vì nội dung có từ này.",
            ],
            [
                "GPA là gì?",
                "Không khớp tốt với văn bản quy chế.",
                "Mở rộng GPA thành điểm trung bình tích lũy/điểm hệ 4/tính điểm trung bình.",
            ],
            [
                "Cách hủy học phần đã đăng ký",
                "Dễ bị kéo sang nguồn tiếng Anh/chứng chỉ hoặc nguồn nhiễu.",
                "Map sang rút bớt học phần, hủy đăng ký học phần, Điều 10/Điều 9.",
            ],
            [
                "F+ và F khác nhau thế nào?",
                "Dễ trả sai hoặc thiếu căn cứ học lại/học đổi.",
                "Kết hợp Điều 16 về thang điểm và Điều 11 về học lại/học đổi.",
            ],
            [
                "Nghỉ học không phép và có phép khác nhau những gì?",
                "Cần ghép nhiều khía cạnh, không chỉ một chunk.",
                "Thêm decomposition rule-first và multi-hop retrieval nhiều nguồn.",
            ],
        ],
    )

    add_heading(doc, "10. Cải tiến prompt và fallback trả lời")
    add_bullets(
        doc,
        [
            "Prompt/context khuyến khích trả lời dựa trên nhiều nguồn khi context có nhiều evidence liên quan.",
            "Nếu context có điều/mục phù hợp, không được dễ dàng trả “không có căn cứ”.",
            "Có fallback extractive khi Gemini lỗi hoặc không khả dụng, tận dụng source đã retrieve được.",
            "Một số câu nghiệp vụ như đăng ký/hủy thi lại có direct answer khi source Web Support rõ ràng.",
            "Câu trả lời cuối vẫn kèm source và sources để frontend/trace hiển thị được nhiều nguồn.",
        ],
    )

    add_heading(doc, "11. Cải tiến trace và khả năng debug")
    add_bullets(
        doc,
        [
            "Ghi route/intent được chọn.",
            "Ghi retrieval debug cho internal, business, website.",
            "Ghi sub-question trong multi-hop.",
            "Ghi danh sách nguồn bị reject/accept và lý do.",
            "Ghi gemini_call_count để kiểm soát quota.",
            "Ghi fallback_decision để biết vì sao rơi sang website hoặc câu trả lời fallback.",
        ],
    )

    add_heading(doc, "12. Bổ sung kiểm thử tự động")
    add_bullets(
        doc,
        [
            "tests/test_gemini_client.py: test xoay vòng key và fallback sang key tiếp theo.",
            "tests/test_business_faq_mapping.py: test không gọi Gemini retrieval plan trong multi-hop sub-question.",
            "tests/test_hybrid_retrieval.py: test query expansion, policy profile và priority nguồn.",
            "tests/test_chat_langchain.py: test aggregate routing, chọn nguồn, deterministic answer và fallback.",
            "tests/test_query_context_routing.py: test nhận diện context sinh viên/chính sách/thao tác.",
        ],
    )

    add_heading(doc, "13. Giá trị đạt được")
    add_bullets(
        doc,
        [
            "Chatbot ổn định hơn khi một Gemini API key hết quota hoặc gặp rate limit.",
            "Giảm chi phí/quota bằng cách hạn chế gọi Gemini trong các bước phụ.",
            "Tăng khả năng tìm đúng tài liệu chính thức, nhất là Quy chế đào tạo đại học chính quy 832.",
            "Giảm nhầm giữa tài liệu đại học, thạc sĩ, tuyển sinh, Web Support và website.",
            "Trả lời tốt hơn với các câu hỏi học vụ dùng ngôn ngữ đời thường.",
            "Có khả năng ghép nhiều nguồn hơn cho câu hỏi phức tạp.",
            "Dễ debug hơn nhờ trace chi tiết và test regression.",
        ],
    )

    add_heading(doc, "14. Hạn chế và hướng phát triển tiếp")
    add_bullets(
        doc,
        [
            "Cơ chế API key luân phiên giúp chống lỗi từng key, nhưng không làm tăng quota nếu tất cả key đều cùng hết hạn mức.",
            "Một số rule học vụ đang được bổ sung theo lỗi thực tế, cần tiếp tục tổng quát hóa để tránh chồng chéo.",
            "Chất lượng trả lời vẫn phụ thuộc vào chất lượng index, metadata và chunking của tài liệu gốc.",
            "Cần chạy bộ ground truth đầy đủ để đo điểm trước/sau một cách định lượng.",
            "Nên theo dõi trace thực tế để phát hiện thêm nhóm câu hỏi còn chọn sai nguồn.",
        ],
    )

    add_heading(doc, "15. Kết luận")
    add_paragraph(
        doc,
        "Từ commit thêm API key luân phiên đến hiện tại, chatbot RAG UNETI đã được cải tiến cả về vận hành "
        "và chất lượng trả lời. Về vận hành, hệ thống hỗ trợ nhiều Gemini API key, xoay vòng key và fallback "
        "khi quota/rate limit. Về chất lượng RAG, hệ thống được bổ sung query expansion, policy profile, scoring "
        "theo nguồn, multi-hop retrieval không gọi thêm Gemini ở bước phụ, fallback an toàn và test regression. "
        "Nhờ đó chatbot chọn nguồn học vụ chính xác hơn, giảm nhầm nguồn và tiết kiệm quota hơn so với trước.",
    )

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = note.add_run("Tài liệu được tạo tự động từ diff Git của dự án.")
    set_run_font(run, size=10, italic=True)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_report()
    print(OUTPUT.resolve())
