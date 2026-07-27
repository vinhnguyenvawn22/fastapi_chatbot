from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "BAO_CAO_LUONG_RAG_TAI_LIEU_CUC_BO_MOI.md"
OUTPUT = ROOT / "BAO_CAO_LUONG_RAG_TAI_LIEU_CUC_BO_MOI.docx"

NAVY = "17365D"
BLUE = "2E75B6"
LIGHT_BLUE = "DCEAF7"
PALE_BLUE = "F3F7FB"
GREEN = "2F7D5B"
GRAY = "5B6573"
LIGHT_GRAY = "E8EBEF"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)


def add_bottom_border(paragraph, color=NAVY, size="12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_run_with_inline_format(paragraph, text: str, code_font=False) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string(NAVY)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
            if code_font:
                run.font.name = "Consolas"
                run.font.size = Pt(9)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("202B38")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for level, size, color in (
        (1, 15, NAVY),
        (2, 12.5, BLUE),
        (3, 11.5, GREEN),
    ):
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(2.5)

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = "FASTAPI CHATBOT  |  LUỒNG RAG TÀI LIỆU CỤC BỘ"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(GRAY)
    add_bottom_border(p, LIGHT_GRAY, "6")

    footer = section.footer
    footer.is_linked_to_previous = False
    add_page_number(footer.paragraphs[0])


def add_cover(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()

    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = label.add_run("BÁO CÁO KỸ THUẬT")
    run.font.name = "Aptos"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("THIẾT KẾ VÀ CẢI TIẾN\nLUỒNG RAG TÀI LIỆU CỤC BỘ")
    run.font.name = "Aptos Display"
    run.font.size = Pt(25)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = rule.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    run.font.color.rgb = RGBColor.from_string(GREEN)
    run.font.size = Pt(12)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(14)
    run = subtitle.add_run("Endpoint /api/chat/local-documents")
    run.font.name = "Consolas"
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    summary = doc.add_table(rows=4, cols=2)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary.autofit = False
    summary.columns[0].width = Cm(4.2)
    summary.columns[1].width = Cm(10.2)
    values = [
        ("Dự án", "FastAPI Chatbot"),
        ("Nguồn dữ liệu", r"C:\fastapi_chatbot\uploads\document"),
        ("Phạm vi", "RAG cục bộ, không tìm kiếm website"),
        ("Ngày báo cáo", "26/07/2026"),
    ]
    for row, (key, value) in zip(summary.rows, values):
        row.cells[0].width = Cm(4.2)
        row.cells[1].width = Cm(10.2)
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            set_cell_margins(cell, 120, 140, 120, 140)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        key_run = row.cells[0].paragraphs[0].add_run(key)
        key_run.bold = True
        key_run.font.color.rgb = RGBColor.from_string(NAVY)
        value_run = row.cells[1].paragraphs[0].add_run(value)
        if key == "Nguồn dữ liệu":
            value_run.font.name = "Consolas"
            value_run.font.size = Pt(9)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "Tách nhu cầu để tìm đúng bằng chứng • Hợp nhất nguồn • Sinh câu trả lời một lần"
    )
    run.font.name = "Aptos"
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string(GRAY)

    doc.add_page_break()


def add_contents(doc: Document) -> None:
    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 1"]
    heading.add_run("Nội dung báo cáo")
    items = [
        "1. Tóm tắt",
        "2. Bối cảnh và vấn đề của luồng cũ",
        "3. Mục tiêu của luồng mới",
        "4. Kiến trúc tổng thể",
        "5. Luồng xử lý chi tiết",
        "6. Ví dụ xử lý",
        "7. Tích hợp với Chat UI",
        "8. Kiểm thử và kết quả",
        "9. Các file chính",
        "10. Ưu điểm của giải pháp",
        "11. Hạn chế còn lại",
        "12. Đề xuất phát triển tiếp",
        "13. Kết luận",
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_page_break()


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    set_cell_margins(cell, 130, 170, 130, 170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string(NAVY)


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    parsed = [
        [cell.strip() for cell in line.strip().strip("|").split("|")]
        for line in lines
    ]
    if len(parsed) >= 2 and all(re.fullmatch(r":?-{3,}:?", c) for c in parsed[1]):
        parsed.pop(1)
    if not parsed:
        return

    table = doc.add_table(rows=len(parsed), cols=len(parsed[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(parsed):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, NAVY)
            elif row_index % 2 == 0:
                set_cell_shading(cell, PALE_BLUE)
            p = cell.paragraphs[0]
            add_run_with_inline_format(p, value)
            for run in p.runs:
                run.font.size = Pt(9)
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F6F8")
    set_cell_margins(cell, 130, 150, 130, 150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(code.rstrip().splitlines()):
        if index:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("263442")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def render_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    skip_metadata = True

    while index < len(lines):
        line = lines[index]

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if skip_metadata:
            if line.startswith("## 1."):
                skip_metadata = False
            else:
                index += 1
                continue

        if not line.strip():
            index += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(doc, table_lines)
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading_match:
            hashes, text = heading_match.groups()
            level = len(hashes) - 1
            p = doc.add_paragraph(style=f"Heading {level}")
            add_run_with_inline_format(p, text)
            if level == 1:
                add_bottom_border(p, LIGHT_BLUE, "8")
            index += 1
            continue

        if line.startswith("> "):
            add_callout(doc, line[2:].strip())
            index += 1
            continue

        bullet_match = re.match(r"^-\s+(.+)$", line)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            add_run_with_inline_format(p, bullet_match.group(1))
            index += 1
            continue

        number_match = re.match(r"^\d+\.\s+(.+)$", line)
        if number_match:
            p = doc.add_paragraph(style="List Number")
            add_run_with_inline_format(p, number_match.group(1))
            index += 1
            continue

        paragraph_lines = [line.strip()]
        index += 1
        while index < len(lines):
            candidate = lines[index]
            if (
                not candidate.strip()
                or candidate.startswith("#")
                or candidate.startswith("|")
                or candidate.startswith("```")
                or candidate.startswith("> ")
                or re.match(r"^-\s+", candidate)
                or re.match(r"^\d+\.\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate.strip())
            index += 1
        p = doc.add_paragraph()
        add_run_with_inline_format(p, " ".join(paragraph_lines))


def build_report() -> Path:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_contents(doc)
    render_markdown(doc, markdown)

    core_properties = doc.core_properties
    core_properties.title = "Báo cáo luồng RAG tài liệu cục bộ mới"
    core_properties.subject = "Endpoint /api/chat/local-documents"
    core_properties.author = "FastAPI Chatbot Project"
    core_properties.keywords = "RAG, FastAPI, BM25, ANN, Gemini, multi-aspect"

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_report())
