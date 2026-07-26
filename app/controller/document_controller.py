from datetime import datetime, timezone
from pathlib import Path
import hashlib
import re
import shutil
import subprocess
import tempfile

from fastapi import HTTPException, UploadFile
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.core.config import (
    CHUNK_OVERLAP,
    CHUNK_SIZE,
    DOCUMENTS_DIR,
    LOCAL_DOCUMENTS_CORPUS,
    LOCAL_DOCUMENTS_INDEX_VERSION,
    MAX_UPLOAD_SIZE_MB,
)
from app.data.query_analyzer import normalize_date, normalize_text


PDF_MIME_TYPES = {"application/pdf", "application/octet-stream"}
SUPPORTED_DOCUMENT_SUFFIXES = {".pdf", ".docx", ".doc", ".xlsx"}
DISCOVERED_DOCUMENT_SUFFIXES = SUPPORTED_DOCUMENT_SUFFIXES
PARSER_VERSION = "local_parser_v1"
CHUNKER_VERSION = "local_chunker_v3"
UNSAFE_FILENAME_PATTERN = re.compile(r'[<>:"\\|?*\x00-\x1f]+')
SECTION_PATTERN = re.compile(
    r"(?im)^\s*((?:Điều|Dieu)\s+\d+[\.\:\s]+.*|(?:Mục|Muc)\s+\d+[\.\:\s]+.*|(?:Chương|Chuong)\s+(?:[IVXLCDM]+|\d+)[\.\:\s]+.*)$"
)
DOCUMENT_TYPE_PATTERN = re.compile(
    r"\b(Quyết định|Quyet dinh|Quy định|Quy dinh|Quy chế|Quy che|Thông báo|Thong bao|Hướng dẫn|Huong dan|Kế hoạch|Ke hoach)\b",
    flags=re.IGNORECASE,
)


def _documents_path() -> Path:
    return Path(DOCUMENTS_DIR).resolve()


def _relative_document_path(file_path: Path) -> str:
    return file_path.resolve().relative_to(_documents_path()).as_posix()


def _extract_source_metadata(file_path: Path) -> dict:
    documents_path = _documents_path()
    relative_path = _relative_document_path(file_path)
    parent_parts = Path(relative_path).parts[:-1]
    phong_ban = None

    for part in reversed(parent_parts):
        normalized = normalize_text(part)
        if normalized.startswith("phong ") or normalized.startswith("trung tam "):
            phong_ban = part
            break

    if not phong_ban and parent_parts:
        phong_ban = parent_parts[0]

    return {
        "phong_ban": phong_ban,
        "relative_path": relative_path,
        "source_root": documents_path.name,
    }


def _safe_pdf_filename(filename: str) -> str:
    name = Path(filename).name.strip()
    name = UNSAFE_FILENAME_PATTERN.sub("_", name)
    name = re.sub(r"\s+", " ", name)

    if not name or name in {".", ".."}:
        raise HTTPException(status_code=400, detail="Ten file khong hop le")

    if not name.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Chi ho tro file PDF")

    return name


def _is_supported_document(file_path: Path) -> bool:
    return file_path.suffix.lower() in SUPPORTED_DOCUMENT_SUFFIXES


def _is_ignored_document(file_path: Path) -> bool:
    """Ignore Office lock files and other transient documents."""
    return file_path.name.startswith("~$")


def _resolve_document_path(file_name: str | Path) -> Path:
    documents_path = _documents_path()
    raw_name = str(file_name or "").replace("\\", "/").strip()

    if not raw_name:
        raise HTTPException(status_code=400, detail="Ten file khong hop le")

    if Path(raw_name).is_absolute():
        file_path = Path(raw_name).resolve()
    else:
        parts = raw_name.split("/")
        if any(part in {"", ".", ".."} for part in parts):
            raise HTTPException(status_code=400, detail="Ten file khong hop le")
        if UNSAFE_FILENAME_PATTERN.search("".join(parts)):
            raise HTTPException(status_code=400, detail="Ten file khong hop le")
        if Path(raw_name).suffix.lower() not in SUPPORTED_DOCUMENT_SUFFIXES:
            raise HTTPException(status_code=400, detail="Chi ho tro PDF, DOCX, DOC va XLSX")
        file_path = (documents_path / raw_name).resolve()

    if documents_path not in file_path.parents:
        raise HTTPException(status_code=400, detail="Ten file khong hop le")

    if not _is_supported_document(file_path):
        raise HTTPException(status_code=400, detail="Chi ho tro PDF, DOCX, DOC va XLSX")

    return file_path


def _file_sha256(file_path: Path) -> str:
    hasher = hashlib.sha256()

    with file_path.open("rb") as file:
        for block in iter(lambda: file.read(1024 * 1024), b""):
            hasher.update(block)

    return hasher.hexdigest()


def _text_sha256(text: str) -> str:
    normalized = " ".join(normalize_text(text).split())
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def _classify_document(file_path: Path, source_metadata: dict) -> dict:
    normalized_name = normalize_text(file_path.name)
    normalized_path = normalize_text(source_metadata["relative_path"])
    suffix = file_path.suffix.lower()
    rag_enabled = normalized_name != "pcntt_mapping_file.docx"
    document_type = "unknown"
    method = "fallback"
    confidence = 0.4

    rules = (
        (("quy che", "quy dinh"), "regulation"),
        (("quyet dinh",), "decision"),
        (("huong dan",), "guideline"),
        (("quy trinh",), "procedure"),
        (("cau hoi", "faq"), "faq"),
        (("bieu mau", "phieu"), "form"),
    )
    for terms, candidate in rules:
        if any(term in normalized_name for term in terms):
            document_type = candidate
            method = "filename_rule"
            confidence = 0.95
            break

    if suffix == ".xlsx" and document_type == "unknown":
        document_type = "spreadsheet"
        method = "file_extension"
        confidence = 0.98
    elif normalized_path.startswith("nghiep_vu/") and document_type == "unknown":
        document_type = "business_document"
        method = "path_rule"
        confidence = 0.9

    return {
        "corpus": LOCAL_DOCUMENTS_CORPUS,
        "index_version": LOCAL_DOCUMENTS_INDEX_VERSION,
        "source_type": "local_file",
        "rag_enabled": rag_enabled,
        "exclude_reason": None if rag_enabled else "legacy_mapping_file",
        "document_type": document_type,
        "department": source_metadata.get("phong_ban"),
        "classification_method": method,
        "classification_confidence": confidence,
        "file_extension": suffix,
        "parser_version": PARSER_VERSION,
        "chunker_version": CHUNKER_VERSION,
    }


def _first_match(pattern: str, text: str, flags=re.IGNORECASE):
    match = re.search(pattern, text, flags=flags)
    return match.group(1).strip() if match else None


def _extract_document_number_from_filename(file_name: str) -> str | None:
    normalized_name = normalize_text(Path(file_name).stem)
    normalized_name = re.sub(r"[_\-.]+", " ", normalized_name)

    patterns = (
        r"\b(?:qd|quyet\s*dinh|tb|thong\s*bao|qc|quy\s*che|vb|van\s*ban|hd|huong\s*dan|qt)\s+(\d{1,6})\b",
        r"\b(\d{1,6})\s+(?:qd|quyet\s*dinh|tb|thong\s*bao|qc|quy\s*che|vb|van\s*ban|hd|huong\s*dan|qt)\b",
    )
    typed_candidates = []
    for pattern in patterns:
        for match in re.finditer(pattern, normalized_name):
            candidate = match.group(1)
            if len(candidate) == 4 and 1900 <= int(candidate) <= 2100:
                continue
            typed_candidates.append(candidate)

    for candidate in typed_candidates:
        if len(candidate) >= 3:
            return candidate

    number_candidates = re.findall(r"\b\d{1,6}\b", normalized_name)
    filtered_candidates = []
    for candidate in number_candidates:
        if len(candidate) == 4 and 1900 <= int(candidate) <= 2100:
            continue
        if len(candidate) == 6:
            continue
        filtered_candidates.append(candidate)

    for candidate in filtered_candidates:
        if len(candidate) >= 3:
            return candidate

    if typed_candidates:
        return typed_candidates[0]

    if filtered_candidates:
        return filtered_candidates[0]

    return None


def _extract_document_metadata(text: str, file_name: str) -> dict:
    lines = [
        " ".join(line.split())
        for line in text.splitlines()
        if line and " ".join(line.split())
    ]
    header_text = "\n".join(lines[:80])
    normalized_header = normalize_text(header_text)
    filename_so_van_ban = _extract_document_number_from_filename(file_name)

    so_van_ban = _first_match(
        r"(?:Số|So)\s*[:\-]?\s*([0-9]{1,6}(?:/[A-Za-z0-9.\-]+)?)",
        header_text,
    )
    if not so_van_ban:
        match = re.search(
            r"(?:so|van\s*ban|quyet\s*dinh|quy\s*dinh|qd)\s*[:\-]?\s*([0-9]{1,6}(?:\s*/\s*[a-z0-9.\-]+)?)",
            normalized_header,
        )
        if not match:
            match = re.search(r"\b([0-9]{2,6})\s*/\s*(?:qd|vb|tb|qc)", normalized_header)
        if match:
            so_van_ban = re.sub(r"\s+", "", match.group(1)).upper()
    if filename_so_van_ban:
        so_van_ban = filename_so_van_ban

    so_van_ban_ngan = None
    if so_van_ban:
        short_match = re.search(r"\d{1,6}", so_van_ban)
        so_van_ban_ngan = short_match.group(0) if short_match else so_van_ban

    ngay_ban_hanh = _first_match(
        r"ngày\s+(\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4})",
        header_text,
    )
    ngay_hieu_luc = _first_match(
        r"hiệu lực(?:\s+thi hành)?(?:\s+kể)?\s+từ\s+ngày\s+(\d{1,2}\s*(?:/|-|\.|tháng)\s*\d{1,2}\s*(?:/|-|\.|năm)?\s*\d{4})",
        header_text,
    )

    type_match = DOCUMENT_TYPE_PATTERN.search(header_text) or DOCUMENT_TYPE_PATTERN.search(file_name)
    loai_van_ban = type_match.group(1) if type_match else None

    don_vi_ban_hanh = None
    for line in lines[:20]:
        normalized_line = normalize_text(line)
        if any(term in normalized_line for term in ("bo ", "truong ", "phong ", "khoa ", "uy ban")):
            don_vi_ban_hanh = line
            break

    ten_van_ban = None
    for line in lines[:80]:
        normalized_line = normalize_text(line)
        if len(line) >= 12 and any(
            term in normalized_line
            for term in ("quy dinh", "quy che", "quyet dinh", "thong bao", "huong dan")
        ):
            ten_van_ban = line
            break
    if not ten_van_ban:
        ten_van_ban = Path(file_name).stem

    return {
        "so_van_ban": so_van_ban,
        "so_van_ban_ngan": so_van_ban_ngan,
        "ngay_ban_hanh": normalize_date(ngay_ban_hanh) if ngay_ban_hanh else None,
        "ngay_hieu_luc": normalize_date(ngay_hieu_luc) if ngay_hieu_luc else None,
        "ten_van_ban": ten_van_ban,
        "don_vi_ban_hanh": don_vi_ban_hanh,
        "loai_van_ban": loai_van_ban,
    }


def _unique_file_path(file_path: Path) -> Path:
    if not file_path.exists():
        return file_path

    stem = file_path.stem
    suffix = file_path.suffix
    parent = file_path.parent

    for index in range(1, 1000):
        candidate = parent / f"{stem}_{index}{suffix}"
        if not candidate.exists():
            return candidate

    raise HTTPException(status_code=409, detail="Khong the tao ten file khong trung")


def list_documents():
    documents_path = _documents_path()

    if not documents_path.exists():
        return []

    files = []

    for file_path in sorted(
        path
        for path in documents_path.rglob("*")
        if (
            path.is_file()
            and path.suffix.lower() in DISCOVERED_DOCUMENT_SUFFIXES
            and not _is_ignored_document(path)
        )
    ):
        stat = file_path.stat()
        source_metadata = _extract_source_metadata(file_path)
        parse_supported = _is_supported_document(file_path)
        files.append({
            "file_name": file_path.name,
            "relative_path": source_metadata["relative_path"],
            "file_path": str(file_path),
            "file_type": file_path.suffix.lower().lstrip("."),
            "parse_supported": parse_supported,
            "index_status": "supported" if parse_supported else "unsupported_file_type",
            "file_size_kb": round(stat.st_size / 1024, 2),
            "updated_at": datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat(),
            "phong_ban": source_metadata["phong_ban"],
            "source_root": source_metadata["source_root"],
        })

    return files


def extract_pdf_text(file_name: str):
    file_path = _resolve_document_path(file_name)

    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"Khong tim thay file: {file_name}")

    try:
        reader = PdfReader(str(file_path))
    except PdfReadError as exc:
        raise HTTPException(status_code=400, detail="File PDF khong doc duoc") from exc
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Khong the mo file PDF") from exc

    text_parts = []

    for page_index, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text_parts.append(f"\n--- Trang {page_index + 1} ---\n{text}")

    return "\n".join(text_parts).strip()


def _extract_docx_path(file_path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="Thieu thu vien python-docx") from exc

    try:
        document = Document(str(file_path))
    except Exception as exc:
        raise HTTPException(status_code=400, detail="File DOCX khong doc duoc") from exc

    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_docx_text(file_name: str):
    return _extract_docx_path(_resolve_document_path(file_name))


def extract_doc_text(file_name: str):
    file_path = _resolve_document_path(file_name)
    converter = shutil.which("soffice") or shutil.which("libreoffice")
    if not converter:
        raise HTTPException(
            status_code=422,
            detail="doc_conversion_failed: LibreOffice/soffice not found",
        )

    with tempfile.TemporaryDirectory(prefix="rag_doc_") as temp_dir:
        result = subprocess.run(
            [
                converter,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                temp_dir,
                str(file_path),
            ],
            capture_output=True,
            text=True,
            timeout=120,
            check=False,
        )
        converted_path = Path(temp_dir) / f"{file_path.stem}.docx"
        if result.returncode != 0 or not converted_path.exists():
            raise HTTPException(status_code=422, detail="doc_conversion_failed")
        return _extract_docx_path(converted_path)


def _markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    padded = [row + [""] * (width - len(row)) for row in rows]
    header = padded[0]
    separator = ["---"] * width

    def render(row):
        return "| " + " | ".join(cell.replace("|", "\\|") for cell in row) + " |"

    return "\n".join([render(header), render(separator), *(render(row) for row in padded[1:])])


def _extract_business_docx_sections(file_path: Path) -> list[dict]:
    """Preserve the module/screen hierarchy encoded in DOCX heading styles."""
    try:
        from docx import Document
        from docx.table import Table
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="Thieu thu vien python-docx") from exc

    try:
        document = Document(str(file_path))
    except Exception as exc:
        raise HTTPException(status_code=400, detail="File DOCX khong doc duoc") from exc

    sections = []
    hierarchy = {}
    current = None
    intro_parts = []

    def flush_current():
        nonlocal current
        if not current:
            return
        body = "\n".join(current.pop("parts")).strip()
        if body:
            current["content"] = body
            sections.append(current)
        current = None

    for block in document.iter_inner_content():
        if isinstance(block, Table):
            rows = [
                [" ".join(cell.text.split()) for cell in row.cells]
                for row in block.rows
                if any(cell.text.strip() for cell in row.cells)
            ]
            table_text = _markdown_table(rows)
            if table_text:
                if current:
                    current["parts"].append(table_text)
                else:
                    intro_parts.append(table_text)
            continue

        text = block.text.strip()
        if not text:
            continue
        style_name = (block.style.name or "").strip()
        heading_match = re.fullmatch(r"Heading\s+(\d+)", style_name, re.IGNORECASE)
        is_title = style_name.lower() == "title"
        if heading_match or is_title:
            flush_current()
            level = int(heading_match.group(1)) if heading_match else 0
            hierarchy = {
                existing_level: value
                for existing_level, value in hierarchy.items()
                if existing_level < level
            }
            hierarchy[level] = text
            section_path = " > ".join(
                hierarchy[key] for key in sorted(hierarchy)
            )
            current = {
                "title": text,
                "heading": text,
                "section_path": section_path,
                "section_type": "business_section",
                "heading_level": level,
                "parts": [],
            }
        elif current:
            current["parts"].append(text)
        else:
            intro_parts.append(text)

    flush_current()
    intro = "\n".join(intro_parts).strip()
    if intro:
        sections.insert(0, {
            "title": "Phan mo dau",
            "heading": "Phan mo dau",
            "section_path": "Phan mo dau",
            "section_type": "business_section",
            "heading_level": 0,
            "content": intro,
        })
    return sections


def _split_business_docx(file_path: Path) -> list[dict]:
    sections = _extract_business_docx_sections(file_path)
    if not sections:
        return _split_general_document(_extract_docx_path(file_path))

    grouped_sections = []
    index = 0
    while index < len(sections):
        section = sections[index]
        normalized_title = normalize_text(section["title"])
        is_step = bool(re.match(r"^(?:buoc|b)\s*\d+\b", normalized_title))
        if not is_step or " > " not in section["section_path"]:
            grouped_sections.append(section)
            index += 1
            continue

        parent_path = section["section_path"].rsplit(" > ", 1)[0]
        step_sections = []
        while index < len(sections):
            candidate = sections[index]
            candidate_title = normalize_text(candidate["title"])
            candidate_parent = (
                candidate["section_path"].rsplit(" > ", 1)[0]
                if " > " in candidate["section_path"]
                else ""
            )
            if (
                not re.match(r"^(?:buoc|b)\s*\d+\b", candidate_title)
                or candidate_parent != parent_path
            ):
                break
            step_sections.append(candidate)
            index += 1

        parent_title = parent_path.rsplit(" > ", 1)[-1]
        grouped_sections.append({
            "title": f"{parent_title} - Cac buoc",
            "heading": parent_title,
            "section_path": parent_path,
            "section_type": "business_workflow",
            "heading_level": section["heading_level"],
            "content": "\n\n".join(
                f'{item["title"]}\n{item["content"]}'
                for item in step_sections
            ),
        })

    sections = grouped_sections
    chunks = []
    for section in sections:
        context = section["section_path"]
        body_chunks = chunk_text(section["content"])
        for split_index, body in enumerate(body_chunks, start=1):
            content = f"{context}\n{body}".strip()
            chunks.append({
                "title": (
                    section["title"]
                    if len(body_chunks) == 1
                    else f'{section["title"]} ({split_index})'
                ),
                "heading": section["heading"],
                "section_path": context,
                "section_type": section["section_type"],
                "heading_level": section["heading_level"],
                "content": content,
                "dieu": None,
                "muc": None,
                "chuong": None,
            })
    return chunks


def extract_xlsx_sections(file_name: str) -> list[dict]:
    file_path = _resolve_document_path(file_name)
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="Thieu thu vien openpyxl") from exc

    try:
        workbook = load_workbook(file_path, read_only=True, data_only=True)
    except Exception as exc:
        raise HTTPException(status_code=400, detail="File XLSX khong doc duoc") from exc

    sections = []
    try:
        for table_index, sheet in enumerate(workbook.worksheets, start=1):
            rows = []
            for raw_row in sheet.iter_rows(values_only=True):
                row = ["" if value is None else " ".join(str(value).split()) for value in raw_row]
                while row and not row[-1]:
                    row.pop()
                if any(row):
                    rows.append(row)
            if not rows:
                continue
            sections.append({
                "title": sheet.title,
                "heading": sheet.title,
                "section_type": "table",
                "table_index": table_index,
                "sheet_name": sheet.title,
                "content": _markdown_table(rows),
                "dieu": None,
                "muc": None,
                "chuong": None,
            })
    finally:
        workbook.close()
    return sections


def extract_document_text(file_name: str):
    file_path = _resolve_document_path(file_name)
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return extract_pdf_text(file_name)
    if suffix == ".docx":
        return extract_docx_text(file_name)
    if suffix == ".doc":
        return extract_doc_text(file_name)
    if suffix == ".xlsx":
        return "\n\n".join(section["content"] for section in extract_xlsx_sections(file_name))

    raise HTTPException(status_code=400, detail=f"Chua ho tro doc file {suffix}")


async def upload_document(file: UploadFile):
    if not file.filename:
        raise HTTPException(status_code=400, detail="Ten file khong hop le")

    if file.content_type and file.content_type not in PDF_MIME_TYPES:
        raise HTTPException(status_code=400, detail="Chi ho tro file PDF")

    safe_name = _safe_pdf_filename(file.filename)
    documents_path = _documents_path()
    documents_path.mkdir(parents=True, exist_ok=True)

    file_path = _unique_file_path((documents_path / safe_name).resolve())
    max_bytes = MAX_UPLOAD_SIZE_MB * 1024 * 1024
    bytes_written = 0

    try:
        with file_path.open("wb") as buffer:
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break

                bytes_written += len(chunk)
                if bytes_written > max_bytes:
                    buffer.close()
                    file_path.unlink(missing_ok=True)
                    raise HTTPException(
                        status_code=413,
                        detail=f"File vuot qua gioi han {MAX_UPLOAD_SIZE_MB}MB",
                    )

                buffer.write(chunk)
    except HTTPException:
        raise
    except Exception as exc:
        file_path.unlink(missing_ok=True)
        raise HTTPException(status_code=500, detail="Khong the luu file upload") from exc
    finally:
        await file.close()

    try:
        PdfReader(str(file_path))
    except Exception as exc:
        file_path.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail="File upload khong phai PDF hop le") from exc

    try:
        from app.data.elasticsearch_client import clear_document_index_cache

        clear_document_index_cache()
    except Exception:
        pass

    vector_index_status = "skipped"
    vector_indexed_chunks = 0

    try:
        from app.data.vector_store import index_chunks

        document_chunks = build_document_chunks(_relative_document_path(file_path))
        vector_indexed_chunks = index_chunks(document_chunks)
        vector_index_status = "indexed"
    except Exception as exc:
        vector_index_status = f"failed: {exc}"

    return {
        "message": "Upload tai lieu thanh cong",
        "file_name": file_path.name,
        "relative_path": _relative_document_path(file_path),
        "file_path": str(file_path),
        "file_size_kb": round(file_path.stat().st_size / 1024, 2),
        "content_hash": _file_sha256(file_path),
        "vector_index_status": vector_index_status,
        "vector_indexed_chunks": vector_indexed_chunks,
    }


def _split_recursive(text: str, chunk_size: int, separators: list[str]) -> list[str]:
    text = text.strip()

    if not text:
        return []

    if len(text) <= chunk_size:
        return [text]

    if not separators:
        return [
            text[start:start + chunk_size].strip()
            for start in range(0, len(text), chunk_size)
            if text[start:start + chunk_size].strip()
        ]

    separator = separators[0]
    raw_parts = text.split(separator)

    if len(raw_parts) == 1:
        return _split_recursive(text, chunk_size, separators[1:])

    parts = [
        f"{part}{separator}" if index < len(raw_parts) - 1 else part
        for index, part in enumerate(raw_parts)
    ]
    chunks = []
    current = ""

    for part in parts:
        part = part.strip()
        if not part:
            continue

        candidate = part if not current else f"{current} {part}"

        if len(candidate) <= chunk_size:
            current = candidate
            continue

        if current:
            chunks.extend(_split_recursive(current, chunk_size, separators[1:]))

        current = part

    if current:
        chunks.extend(_split_recursive(current, chunk_size, separators[1:]))

    return chunks


def _with_overlap(chunks: list[str], overlap: int, chunk_size: int) -> list[str]:
    if overlap <= 0 or len(chunks) <= 1:
        return chunks

    overlapped = [chunks[0]]

    for previous, current in zip(chunks, chunks[1:]):
        prefix = previous[-overlap:].strip()
        if len(previous) > overlap and " " in prefix:
            prefix = prefix.split(" ", 1)[1].strip()

        combined = f"{prefix} {current}".strip()

        if len(combined) > chunk_size + overlap:
            combined = combined[-(chunk_size + overlap):].strip()

        overlapped.append(combined)

    return overlapped


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP):
    if chunk_size <= 0:
        raise ValueError("chunk_size phai lon hon 0")

    if overlap < 0:
        raise ValueError("overlap khong duoc am")

    if overlap >= chunk_size:
        overlap = max(chunk_size // 5, 0)

    separators = ["\n\n", "\n", ". ", "; ", ", ", " "]
    chunks = _split_recursive(text, chunk_size, separators)

    return _with_overlap(chunks, overlap, chunk_size)


def split_text_by_metadata(text: str):
    matches = list(SECTION_PATTERN.finditer(text))

    if not matches:
        fallback_chunks = chunk_text(text)
        return [
            {
                "title": f"Đoạn {index}",
                "dieu": None,
                "muc": None,
                "chuong": None,
                "content": chunk,
            }
            for index, chunk in enumerate(fallback_chunks, start=1)
        ]

    chunks = []
    current_chuong = None
    current_muc = None

    if matches[0].start() > 0:
        intro = text[:matches[0].start()].strip()
        if intro:
            for index, chunk in enumerate(chunk_text(intro), start=1):
                title = "Phần mở đầu" if index == 1 else f"Phần mở đầu ({index})"
                chunks.append({
                    "title": title,
                    "dieu": None,
                    "muc": None,
                    "chuong": None,
                    "content": chunk,
                })

    for index, match in enumerate(matches):
        title = match.group(1).strip()
        start = match.start()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        content = text[start:end].strip()

        chuong_match = re.search(r"(?:Chương|Chuong)\s+([IVXLCDM]+|\d+)", title, flags=re.IGNORECASE)
        if chuong_match:
            current_chuong = chuong_match.group(1).upper()
            current_muc = None

        muc_match = re.search(r"(?:Mục|Muc)\s+(\d+)", title, flags=re.IGNORECASE)
        if muc_match:
            current_muc = int(muc_match.group(1))

        dieu_match = re.search(r"(?:Điều|Dieu)\s+(\d+)", title, flags=re.IGNORECASE)
        dieu = int(dieu_match.group(1)) if dieu_match else None

        split_chunks = chunk_text(content)
        for split_index, split_chunk in enumerate(split_chunks, start=1):
            chunk_title = title if len(split_chunks) == 1 else f"{title} ({split_index})"
            chunks.append({
                "title": chunk_title,
                "dieu": dieu,
                "muc": current_muc,
                "chuong": current_chuong,
                "content": split_chunk,
            })

    return chunks


def _split_general_document(text: str) -> list[dict]:
    heading_pattern = re.compile(
        r"(?m)^\s*((?:\d+(?:\.\d+)*[\.\)]\s+|Bước\s+\d+[\.\:\s]+|Câu\s+\d+[\.\:\s]+).+)$",
        flags=re.IGNORECASE,
    )
    matches = list(heading_pattern.finditer(text))
    if not matches:
        return [
            {
                "title": f"Doan {index}",
                "heading": None,
                "section_type": "content",
                "content": content,
                "dieu": None,
                "muc": None,
                "chuong": None,
            }
            for index, content in enumerate(chunk_text(text), start=1)
        ]

    sections = []
    if matches[0].start() > 0:
        intro = text[:matches[0].start()].strip()
        for content in chunk_text(intro):
            sections.append({
                "title": "Phan mo dau",
                "heading": "Phan mo dau",
                "section_type": "content",
                "content": content,
                "dieu": None,
                "muc": None,
                "chuong": None,
            })

    for index, match in enumerate(matches):
        heading = match.group(1).strip()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        section_text = text[match.start():end].strip()
        for split_index, content in enumerate(chunk_text(section_text), start=1):
            sections.append({
                "title": heading if split_index == 1 else f"{heading} ({split_index})",
                "heading": heading,
                "section_type": "content",
                "content": content,
                "dieu": None,
                "muc": None,
                "chuong": None,
            })
    return sections


def build_document_chunks(file_name: str):
    file_path = _resolve_document_path(file_name)
    text = extract_document_text(file_name)
    source_metadata = _extract_source_metadata(file_path)
    classification = _classify_document(file_path, source_metadata)
    if file_path.suffix.lower() == ".xlsx":
        raw_sections = extract_xlsx_sections(file_name)
        chunks = []
        for section in raw_sections:
            split_contents = chunk_text(section["content"])
            for split_index, content in enumerate(split_contents, start=1):
                chunks.append({
                    **section,
                    "title": (
                        section["title"]
                        if len(split_contents) == 1
                        else f'{section["title"]} ({split_index})'
                    ),
                    "content": content,
                })
    elif (
        classification["document_type"] == "business_document"
        and file_path.suffix.lower() == ".docx"
    ):
        chunks = _split_business_docx(file_path)
    elif classification["document_type"] in {"regulation", "decision"}:
        chunks = split_text_by_metadata(text)
    else:
        chunks = _split_general_document(text)

    document_metadata = _extract_document_metadata(text, file_path.name)
    stat = file_path.stat()
    file_hash = _file_sha256(file_path)
    content_hash = _text_sha256(text)
    document_id = content_hash[:24]

    documents = []
    seen_chunk_hashes = set()

    for index, chunk in enumerate(chunks, start=1):
        chunk_hash = _text_sha256(chunk["content"])
        if chunk_hash in seen_chunk_hashes:
            continue
        seen_chunk_hashes.add(chunk_hash)
        section_parts = [
            f"Chuong {chunk.get('chuong')}" if chunk.get("chuong") else None,
            f"Muc {chunk.get('muc')}" if chunk.get("muc") is not None else None,
            f"Dieu {chunk.get('dieu')}" if chunk.get("dieu") is not None else None,
            chunk.get("heading"),
        ]
        section_path = (
            chunk.get("section_path")
            or " > ".join(part for part in section_parts if part)
        )
        documents.append({
            "doc_name": file_path.name,
            "relative_path": source_metadata["relative_path"],
            "phong_ban": source_metadata["phong_ban"],
            "source_root": source_metadata["source_root"],
            "title": chunk["title"],
            "dieu": chunk["dieu"],
            "muc": chunk.get("muc"),
            "chuong": chunk.get("chuong"),
            "heading": chunk.get("heading"),
            "section_type": chunk.get("section_type", "content"),
            "heading_level": chunk.get("heading_level"),
            "table_index": chunk.get("table_index"),
            "sheet_name": chunk.get("sheet_name"),
            "section_path": section_path or chunk.get("title"),
            "parent_section_id": _text_sha256(section_path or chunk.get("title", ""))[:24],
            "chunk_index": index,
            "content": chunk["content"],
            "file_path": str(file_path),
            "is_active": True,
            "file_hash": file_hash,
            "content_hash": content_hash,
            "chunk_hash": chunk_hash,
            "document_id": document_id,
            "updated_at": datetime.fromtimestamp(stat.st_mtime, timezone.utc).isoformat(),
            **classification,
            **document_metadata,
        })

    return documents
