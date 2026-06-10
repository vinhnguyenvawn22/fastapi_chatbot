from collections import OrderedDict
from collections.abc import Mapping
from datetime import datetime
from html import unescape
from io import BytesIO
from urllib.parse import quote_plus, urljoin, urlparse
from urllib.request import Request, urlopen
from urllib.error import URLError
from xml.etree import ElementTree
import zipfile
import hashlib
import json
import math
import re
import ssl
import time

from google.protobuf.json_format import MessageToDict
from pypdf import PdfReader

from app.core.config import (
    DISCOVERY_COLLECTION_ID,
    DISCOVERY_ENGINE_ID,
    DISCOVERY_LOCATION,
    DISCOVERY_PROJECT_NUMBER,
    DISCOVERY_SERVING_CONFIG_ID,
    RETRIEVAL_CACHE_MAX_ITEMS,
    RETRIEVAL_CACHE_TTL_SECONDS,
    UNETI_WEBSITE_DOMAIN,
    WEBSITE_EXTRACT_MAX_CHARS,
    WEBSITE_FETCH_MAX_BYTES,
    WEBSITE_FETCH_TIMEOUT,
    WEBSITE_MIN_SOURCE_SCORE,
    WEBSITE_RERANK_TOP_K,
    WEBSITE_SEARCH_TOP_K,
)
from app.controller.document_controller import chunk_text
from app.data.elasticsearch_client import clear_document_index_cache, get_keywords, normalize_text
from app.data.gemini_client import ask_gemini
from app.data.vector_store import index_chunks


NO_WEBSITE_CONFIG_ANSWER = (
    "Chua cau hinh Google Vertex AI Search/Discovery Engine cho website UNETI."
)
NO_WEBSITE_RESULTS_ANSWER = "Chua tim thay thong tin phu hop tren website UNETI."
USER_AGENT = "UNETI-RAG-Assistant/1.0"
ATTACHMENT_EXTENSIONS = (".pdf", ".docx", ".doc")
WEBSITE_SOURCE_ROOT = "UNETI website"
FALLBACK_SITEMAP_LIMIT = 200
FALLBACK_CATEGORY_PATHS = (
    "/category/dao-tao/thong-bao-ke-hoach-dao-tao/",
)
_WEBSITE_INDEX_CACHE = OrderedDict()


def _validate_config():
    missing = []
    if not DISCOVERY_PROJECT_NUMBER:
        missing.append("DISCOVERY_PROJECT_NUMBER")
    if not DISCOVERY_ENGINE_ID:
        missing.append("DISCOVERY_ENGINE_ID")
    return missing


def _is_uneti_url(url: str) -> bool:
    host = urlparse(url or "").netloc.lower()
    domain = UNETI_WEBSITE_DOMAIN.lower()
    return host == domain or host.endswith(f".{domain}")


def _struct_to_dict(value) -> dict:
    if not value:
        return {}

    if isinstance(value, Mapping):
        return dict(value)

    try:
        return MessageToDict(value)
    except Exception:
        return {}


def _first_text(value) -> str:
    if isinstance(value, str):
        return value
    if hasattr(value, "string_value") and value.string_value:
        return value.string_value
    if hasattr(value, "list_value"):
        return _first_text(value.list_value.values)
    if hasattr(value, "struct_value"):
        return _first_text(value.struct_value.fields)
    if isinstance(value, Mapping):
        return _first_text(dict(value))
    if isinstance(value, list):
        for item in value:
            text = _first_text(item)
            if text:
                return text
    if isinstance(value, dict):
        for key in ("snippet", "content", "text", "stringValue"):
            text = _first_text(value.get(key))
            if text:
                return text
        for item in value.values():
            text = _first_text(item)
            if text:
                return text
    return ""


def _extract_document(result) -> dict:
    document = getattr(result, "document", None)
    if not document:
        return {}

    data = _struct_to_dict(getattr(document, "derived_struct_data", None))
    if not data:
        data = _struct_to_dict(getattr(document, "derivedStructData", None))

    title = _first_text(data.get("title")) or getattr(document, "name", "") or "Khong tieu de"
    url = _first_text(data.get("link")) or _first_text(data.get("uri")) or _first_text(data.get("url"))
    snippet = (
        _first_text(data.get("snippets"))
        or _first_text(data.get("extractive_answers"))
        or _first_text(data.get("description"))
    )

    return {
        "title": title,
        "url": url,
        "snippet": snippet,
    }


def _score_source(question: str, source: dict) -> float:
    query_terms = set(get_keywords(question))
    searchable = normalize_text(
        " ".join([
            source.get("title", ""),
            source.get("snippet", ""),
            source.get("url", ""),
        ])
    )
    source_terms = set(get_keywords(searchable))

    overlap = len(query_terms & source_terms)
    score = overlap * 10.0

    normalized_question = normalize_text(question)
    if normalized_question and normalized_question in searchable:
        score += 25.0

    query_years = set(re.findall(r"\b20\d{2}\b", question))
    source_years = set(re.findall(r"\b20\d{2}\b", searchable))
    if query_years:
        if query_years & source_years:
            score += 45.0
        elif source_years:
            score -= 45.0

    title = normalize_text(source.get("title", ""))
    score += sum(4.0 for term in query_terms if term in title)

    for phrase in ("tuyen sinh", "thong bao", "dao tao", "hoc bong", "viec lam"):
        if phrase in normalize_text(question):
            if phrase in title:
                score += 30.0
            if phrase in searchable:
                score += 15.0

    url = source.get("url", "")
    if _is_uneti_url(url):
        score += 20.0

    snippet_length = len(source.get("snippet", ""))
    if snippet_length:
        score += min(math.log(snippet_length + 1), 8.0)

    return round(score, 4)


def _fallback_source_matches_query(question: str, source: dict) -> bool:
    query_terms = set(get_keywords(question))
    searchable = normalize_text(
        " ".join([
            source.get("title", ""),
            source.get("snippet", ""),
            source.get("url", ""),
        ])
    )
    source_terms = set(get_keywords(searchable))
    query_years = set(re.findall(r"\b20\d{2}\b", question))
    source_years = set(re.findall(r"\b20\d{2}\b", searchable))

    if query_years and not (query_years & source_years):
        return False

    non_year_terms = {term for term in query_terms if not re.fullmatch(r"20\d{2}", term)}
    overlap = non_year_terms & source_terms

    if len(non_year_terms) >= 4:
        return len(overlap) >= 2

    return bool(overlap)


def _dedupe_and_rerank(question: str, raw_sources: list[dict]) -> list[dict]:
    by_url = {}

    for source in raw_sources:
        url = source.get("url") or ""
        if not url or not _is_uneti_url(url):
            continue

        key = url.split("#", 1)[0].rstrip("/")
        source["score"] = _score_source(question, source)

        if key not in by_url or source["score"] > by_url[key]["score"]:
            by_url[key] = source

    ranked = sorted(by_url.values(), key=lambda item: item["score"], reverse=True)
    return [
        source
        for source in ranked
        if source["score"] >= WEBSITE_MIN_SOURCE_SCORE
    ][:WEBSITE_RERANK_TOP_K]


def _dedupe_and_rerank_fallback(question: str, raw_sources: list[dict]) -> list[dict]:
    by_url = {}

    for source in raw_sources:
        url = source.get("url") or ""
        if not url or not _is_uneti_url(url):
            continue

        if not _fallback_source_matches_query(question, source):
            continue

        key = url.split("#", 1)[0].rstrip("/")
        source["score"] = _score_source(question, source)

        if key not in by_url or source["score"] > by_url[key]["score"]:
            by_url[key] = source

    ranked = sorted(by_url.values(), key=lambda item: item["score"], reverse=True)
    return [
        source
        for source in ranked
        if source["score"] > 0
    ][:WEBSITE_RERANK_TOP_K]


def _read_response(response) -> tuple[str, bytes]:
    content_type = response.headers.get("content-type", "")
    content_length = response.headers.get("content-length")
    if content_length and int(content_length) > WEBSITE_FETCH_MAX_BYTES:
        raise ValueError("Remote file is larger than WEBSITE_FETCH_MAX_BYTES")

    chunks = []
    total = 0
    while True:
        chunk = response.read(64 * 1024)
        if not chunk:
            break
        total += len(chunk)
        if total > WEBSITE_FETCH_MAX_BYTES:
            raise ValueError("Remote file is larger than WEBSITE_FETCH_MAX_BYTES")
        chunks.append(chunk)

    return content_type.lower(), b"".join(chunks)


def _fetch_url(url: str) -> tuple[str, bytes]:
    if not _is_uneti_url(url):
        raise ValueError("Blocked non-UNETI URL")

    request = Request(url, headers={"User-Agent": USER_AGENT})
    try:
        with urlopen(request, timeout=WEBSITE_FETCH_TIMEOUT) as response:
            return _read_response(response)
    except URLError as exc:
        if "CERTIFICATE_VERIFY_FAILED" not in str(exc):
            raise

    # UNETI can serve a certificate chain that Python on Windows may not verify.
    # The retry is still domain-limited by _is_uneti_url above.
    insecure_context = ssl._create_unverified_context()
    with urlopen(request, timeout=WEBSITE_FETCH_TIMEOUT, context=insecure_context) as response:
        return _read_response(response)


def _title_from_url(url: str) -> str:
    path = urlparse(url or "").path.strip("/")
    if not path:
        return UNETI_WEBSITE_DOMAIN

    slug = path.rsplit("/", 1)[-1]
    title = re.sub(r"[-_]+", " ", unescape(slug)).strip()
    return title or path


def _search_wordpress_rest(question: str) -> list[dict]:
    search_url = (
        f"https://{UNETI_WEBSITE_DOMAIN}/wp-json/wp/v2/search"
        f"?search={quote_plus(question)}&per_page={WEBSITE_SEARCH_TOP_K}"
    )
    content_type, data = _fetch_url(search_url)
    if "json" not in content_type:
        return []

    try:
        payload = json.loads(data.decode("utf-8"))
    except Exception:
        return []

    sources = []
    for item in payload if isinstance(payload, list) else []:
        url = item.get("url") or item.get("link") or ""
        if not _is_uneti_url(url):
            continue

        title = _first_text(item.get("title")) or _title_from_url(url)
        subtype = item.get("subtype") or item.get("type") or ""
        sources.append({
            "title": title,
            "url": url,
            "snippet": subtype,
        })

    return sources


def _search_site_html(question: str) -> list[dict]:
    search_url = f"https://{UNETI_WEBSITE_DOMAIN}/?s={quote_plus(question)}"
    content_type, data = _fetch_url(search_url)
    html_text = _decode_html(data, content_type)
    link_pattern = re.compile(
        r"""<a\b[^>]*href\s*=\s*["']([^"']+)["'][^>]*>(.*?)</a>""",
        re.IGNORECASE | re.DOTALL,
    )
    sources = []
    seen = set()

    for href, raw_title in link_pattern.findall(html_text):
        url = urljoin(search_url, unescape(href).strip())
        if not _is_uneti_url(url):
            continue

        parsed = urlparse(url)
        if parsed.fragment or parsed.scheme not in {"http", "https"}:
            continue

        path = parsed.path.strip("/")
        if not path or any(part in path for part in ("wp-content", "wp-json", "xmlrpc.php")):
            continue

        key = url.split("#", 1)[0].rstrip("/")
        if key in seen:
            continue
        seen.add(key)

        title = re.sub(r"(?is)<[^>]+>", " ", raw_title)
        title = " ".join(unescape(title).split()) or _title_from_url(url)
        sources.append({
            "title": title,
            "url": key,
            "snippet": title,
        })

    return sources


def _search_category_pages(question: str) -> list[dict]:
    sources = []
    seen = set()

    for path in FALLBACK_CATEGORY_PATHS:
        for page_index in range(1, 4):
            category_url = f"https://{UNETI_WEBSITE_DOMAIN}{path}"
            if page_index > 1:
                category_url = urljoin(category_url, f"page/{page_index}/")

            try:
                content_type, data = _fetch_url(category_url)
            except Exception:
                continue

            html_text = _decode_html(data, content_type)
            link_pattern = re.compile(
                r"""<a\b[^>]*href\s*=\s*["']([^"']+)["'][^>]*>(.*?)</a>""",
                re.IGNORECASE | re.DOTALL,
            )

            for href, raw_title in link_pattern.findall(html_text):
                url = urljoin(category_url, unescape(href).strip())
                if not _is_uneti_url(url):
                    continue

                parsed = urlparse(url)
                if parsed.fragment or parsed.scheme not in {"http", "https"}:
                    continue

                path = parsed.path.strip("/")
                if not path or any(part in path for part in ("category/", "wp-content", "wp-json")):
                    continue

                key = url.split("#", 1)[0].rstrip("/")
                if key in seen:
                    continue
                seen.add(key)

                title = re.sub(r"(?is)<[^>]+>", " ", raw_title)
                title = " ".join(unescape(title).split()) or _title_from_url(url)
                sources.append({
                    "title": title,
                    "url": key,
                    "snippet": title,
                })

    return sources


def _extract_sitemap_locations(data: bytes, content_type: str = "") -> tuple[list[str], list[str]]:
    xml_text = _decode_html(data, content_type)

    try:
        root = ElementTree.fromstring(xml_text)
    except ElementTree.ParseError:
        return [], []

    tag_name = root.tag.rsplit("}", 1)[-1].lower()
    locations = [
        (node.text or "").strip()
        for node in root.iter()
        if node.tag.rsplit("}", 1)[-1].lower() == "loc" and (node.text or "").strip()
    ]

    if tag_name == "sitemapindex":
        return [], locations

    return locations, []


def _search_sitemap(question: str) -> list[dict]:
    sitemap_queue = [
        f"https://{UNETI_WEBSITE_DOMAIN}/sitemap.xml",
        f"https://{UNETI_WEBSITE_DOMAIN}/post-sitemap.xml",
        f"https://{UNETI_WEBSITE_DOMAIN}/page-sitemap.xml",
    ]
    seen_sitemaps = set()
    page_urls = []

    while sitemap_queue and len(page_urls) < FALLBACK_SITEMAP_LIMIT:
        sitemap_url = sitemap_queue.pop(0)
        if sitemap_url in seen_sitemaps or not _is_uneti_url(sitemap_url):
            continue

        seen_sitemaps.add(sitemap_url)

        try:
            content_type, data = _fetch_url(sitemap_url)
            urls, child_sitemaps = _extract_sitemap_locations(data, content_type)
        except Exception:
            continue

        page_urls.extend(url for url in urls if _is_uneti_url(url))

        for child_url in child_sitemaps:
            if len(seen_sitemaps) + len(sitemap_queue) >= 8:
                break
            if _is_uneti_url(child_url):
                sitemap_queue.append(child_url)

    query_terms = set(get_keywords(question))
    sources = []
    for url in page_urls[:FALLBACK_SITEMAP_LIMIT]:
        title = _title_from_url(url)
        searchable = set(get_keywords(f"{title} {url}"))
        if query_terms and not (query_terms & searchable):
            continue

        sources.append({
            "title": title,
            "url": url,
            "snippet": title,
        })

    return sources


def _decode_html(data: bytes, content_type: str = "") -> str:
    charset_match = re.search(r"charset=([\w-]+)", content_type or "", re.IGNORECASE)
    encodings = [charset_match.group(1)] if charset_match else []
    encodings.extend(["utf-8", "utf-8-sig", "cp1258", "latin-1"])

    for encoding in encodings:
        try:
            return data.decode(encoding)
        except (LookupError, UnicodeDecodeError):
            continue
    return data.decode("utf-8", errors="ignore")


def _html_to_text(html_text: str) -> str:
    text = re.sub(r"(?is)<(script|style|noscript).*?</\1>", " ", html_text)
    text = re.sub(r"(?is)<br\s*/?>", "\n", text)
    text = re.sub(r"(?is)</(p|div|li|tr|h[1-6])>", "\n", text)
    text = re.sub(r"(?is)<[^>]+>", " ", text)
    text = unescape(text)
    lines = [" ".join(line.split()) for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def _looks_like_attachment(url: str) -> bool:
    parsed = urlparse(url)
    path = parsed.path.lower()
    query = parsed.query.lower()

    if path.endswith(ATTACHMENT_EXTENSIONS):
        return True

    return any(marker in path or marker in query for marker in (
        "download",
        "wpdmdl",
        "attachment",
    ))


def _extract_attachment_links(page_url: str, html_text: str) -> list[str]:
    candidates = []
    attr_pattern = re.compile(r"""(?:href|src|data)\s*=\s*["']([^"']+)["']""", re.IGNORECASE)

    for match in attr_pattern.finditer(html_text):
        candidate = urljoin(page_url, unescape(match.group(1).strip()))
        if _is_uneti_url(candidate) and _looks_like_attachment(candidate):
            candidates.append(candidate)

    deduped = []
    seen = set()
    for candidate in candidates:
        key = candidate.split("#", 1)[0]
        if key in seen:
            continue
        seen.add(key)
        deduped.append(candidate)

    return deduped


def _extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    parts = []

    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())

    return "\n\n".join(parts)


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except Exception:
        return _extract_docx_text_with_stdlib(data)

    document = Document(BytesIO(data))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))

    return "\n".join(parts)


def _extract_docx_text_with_stdlib(data: bytes) -> str:
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    parts = []

    with zipfile.ZipFile(BytesIO(data)) as archive:
        document_xml = archive.read("word/document.xml")

    root = ElementTree.fromstring(document_xml)
    for paragraph in root.findall(".//w:p", namespaces):
        texts = [
            node.text
            for node in paragraph.findall(".//w:t", namespaces)
            if node.text
        ]
        line = "".join(texts).strip()
        if line:
            parts.append(line)

    return "\n".join(parts)


def _extract_attachment_content(url: str) -> dict:
    content_type, data = _fetch_url(url)
    path = urlparse(url).path.lower()

    if path.endswith(".pdf") or "pdf" in content_type:
        text = _extract_pdf_text(data)
        if text.strip():
            return {"text": text, "scan_skipped": False}

        return {"text": "", "scan_skipped": True}

    if path.endswith(".docx") or "officedocument.wordprocessingml.document" in content_type:
        return {"text": _extract_docx_text(data), "scan_skipped": False}

    if path.endswith(".doc"):
        return {"text": "", "scan_skipped": False}

    if "html" in content_type:
        return {"text": _html_to_text(_decode_html(data, content_type)), "scan_skipped": False}

    return {"text": "", "scan_skipped": False}


def _extract_attachment_text(url: str) -> str:
    return _extract_attachment_content(url).get("text", "")


def _enrich_source_content(source: dict) -> dict:
    enriched = dict(source)
    url = source.get("url") or ""
    fetch_debug = {
        "url": url,
        "attachment_url": None,
        "content_chars": 0,
        "error": None,
    }

    try:
        if _looks_like_attachment(url):
            attachment = _extract_attachment_content(url)
            text = attachment.get("text", "")
            enriched["attachment_url"] = url
            enriched["content"] = text[:WEBSITE_EXTRACT_MAX_CHARS]
            enriched["scan_skipped"] = bool(attachment.get("scan_skipped"))
            fetch_debug["attachment_url"] = url
            fetch_debug["content_chars"] = len(enriched["content"])
            fetch_debug["scan_skipped"] = enriched["scan_skipped"]
            enriched["fetch_debug"] = fetch_debug
            return enriched

        content_type, data = _fetch_url(url)
        html_text = _decode_html(data, content_type)
        page_text = _html_to_text(html_text)
        attachment_links = _extract_attachment_links(url, html_text)
        fetch_debug["attachment_candidates"] = attachment_links[:5]

        for attachment_url in attachment_links[:5]:
            try:
                attachment = _extract_attachment_content(attachment_url)
                attachment_text = attachment.get("text", "")
            except Exception as exc:
                fetch_debug.setdefault("attachment_errors", []).append({
                    "url": attachment_url,
                    "error": str(exc),
                })
                continue

            if attachment_text.strip():
                enriched["attachment_url"] = attachment_url
                enriched["content"] = attachment_text[:WEBSITE_EXTRACT_MAX_CHARS]
                enriched["scan_skipped"] = False
                fetch_debug["attachment_url"] = attachment_url
                fetch_debug["content_chars"] = len(enriched["content"])
                enriched["fetch_debug"] = fetch_debug
                return enriched

            if attachment.get("scan_skipped"):
                fetch_debug.setdefault("scan_skipped_attachments", []).append(attachment_url)

        if attachment_links:
            skipped_attachment_url = None
            skipped_attachments = fetch_debug.get("scan_skipped_attachments") or []
            if skipped_attachments:
                skipped_attachment_url = skipped_attachments[0]

            combined_text = "\n\n".join(
                part for part in (source.get("snippet", ""), page_text) if str(part).strip()
            )
            enriched["attachment_url"] = skipped_attachment_url
            enriched["content"] = combined_text[:WEBSITE_EXTRACT_MAX_CHARS]
            enriched["scan_skipped"] = bool(fetch_debug.get("scan_skipped_attachments"))
            fetch_debug["attachment_url"] = skipped_attachment_url
            fetch_debug["content_chars"] = len(enriched["content"])
            enriched["fetch_debug"] = fetch_debug
            return enriched

        combined_text = "\n\n".join(
            part for part in (source.get("snippet", ""), page_text) if str(part).strip()
        )
        enriched["content"] = combined_text[:WEBSITE_EXTRACT_MAX_CHARS]
        fetch_debug["content_chars"] = len(enriched["content"])
    except Exception as exc:
        enriched["content"] = source.get("snippet", "")
        fetch_debug["error"] = str(exc)
        fetch_debug["content_chars"] = len(enriched["content"])

    enriched["fetch_debug"] = fetch_debug
    return enriched


def _answerable_sources_from_selected(selected_sources: list[dict]) -> tuple[list[dict], list[dict], bool]:
    enriched_sources = [_enrich_source_content(source) for source in selected_sources]
    top_source_scan_skipped = bool(
        enriched_sources
        and enriched_sources[0].get("scan_skipped")
        and not (enriched_sources[0].get("content") or "").strip()
    )
    answerable_sources = [] if top_source_scan_skipped else [
        source
        for source in enriched_sources
        if (source.get("content") or "").strip()
    ]

    return answerable_sources, enriched_sources, top_source_scan_skipped


def _fallback_search_and_extract_website_sources(
    question: str,
    debug: dict | None = None,
    reason: str = "primary_search_unavailable",
) -> dict:
    raw_sources = []
    fallback_errors = []

    try:
        raw_sources.extend(_search_wordpress_rest(question))
    except Exception as exc:
        fallback_errors.append({"mode": "wordpress_rest", "error": str(exc)})

    if not raw_sources:
        try:
            raw_sources.extend(_search_category_pages(question))
        except Exception as exc:
            fallback_errors.append({"mode": "category_pages", "error": str(exc)})

    if not raw_sources:
        try:
            raw_sources.extend(_search_site_html(question))
        except Exception as exc:
            fallback_errors.append({"mode": "site_html_search", "error": str(exc)})

    if not raw_sources:
        try:
            raw_sources.extend(_search_sitemap(question))
        except Exception as exc:
            fallback_errors.append({"mode": "sitemap", "error": str(exc)})

    selected_sources = _dedupe_and_rerank_fallback(question, raw_sources)
    answerable_sources, enriched_sources, top_source_scan_skipped = _answerable_sources_from_selected(
        selected_sources
    )

    if debug is not None:
        debug.update({
            "status": "fallback_success" if answerable_sources else "fallback_no_answerable_sources",
            "fallback_reason": reason,
            "fallback_errors": fallback_errors,
            "raw_results_count": len(raw_sources),
            "selected_sources_count": len(selected_sources),
            "answerable_sources_count": len(answerable_sources),
            "mode": "wordpress_rest_or_category_or_site_html_or_sitemap",
            "top_source_scan_skipped": top_source_scan_skipped,
            "selected_sources": [
                {
                    "title": source.get("title"),
                    "url": source.get("url"),
                    "attachment_url": source.get("attachment_url"),
                    "score": source.get("score"),
                    "snippet_chars": len(source.get("snippet") or ""),
                    "content_chars": len(source.get("content") or ""),
                    "scan_skipped": bool(source.get("scan_skipped")),
                    "fetch_debug": source.get("fetch_debug"),
                }
                for source in enriched_sources
            ],
        })

    return {"sources": answerable_sources}


def _build_website_prompt(question: str, sources: list[dict]) -> str:
    now = datetime.now().strftime("%d/%m/%Y %H:%M")
    context = "\n\n".join(
        [
            (
                f'<NGUON_WEB id="{index}" title="{source["title"]}" url="{source["url"]}" '
                f'attachment_url="{source.get("attachment_url", "")}" score="{source["score"]}">\n'
                f'{source.get("content") or source.get("snippet", "")}\n</NGUON_WEB>'
            )
            for index, source in enumerate(sources, start=1)
        ]
    )

    return f"""
Ban la tro ly tra cuu thong tin tren website chinh thuc cua Truong Dai hoc Kinh te - Ky thuat Cong nghiep.

THOI GIAN HE THONG: {now}

CHI DAN:
- Chi tra loi dua tren NGUON_WEB ben duoi.
- Khong tu bia thong tin ngoai nguon.
- Neu nguon chua du de tra loi, noi rang chua tim thay thong tin phu hop tren website UNETI.
- Tra loi ngan gon, di thang vao noi dung.
- Cuoi cau tra loi phai co dong nguon dang: (Nguon: [title] - [url])

NGUON_WEB:
{context}

CAU HOI:
{question}

TRA LOI:
""".strip()


def _website_content_hash(source: dict) -> str:
    identity = "|".join([
        source.get("url", ""),
        source.get("attachment_url", ""),
        source.get("title", ""),
    ])
    return hashlib.sha256(identity.encode("utf-8", errors="ignore")).hexdigest()


def _doc_name_from_url(url: str) -> str:
    path = urlparse(url or "").path
    name = path.rsplit("/", 1)[-1]
    return unescape(name) or "website_uneti"


def _build_website_chunks(source: dict) -> list[dict]:
    content = (source.get("content") or "").strip()
    if not content:
        return []

    content_hash = _website_content_hash(source)
    doc_name = _doc_name_from_url(source.get("attachment_url") or source.get("url"))
    relative_path = source.get("attachment_url") or source.get("url")
    chunks = chunk_text(content)
    now = datetime.now().isoformat()

    return [
        {
            "doc_name": doc_name,
            "relative_path": relative_path,
            "phong_ban": source.get("phong_ban") or source.get("source_root") or "Website",
            "source_root": source.get("source_root") or WEBSITE_SOURCE_ROOT,
            "title": source.get("title") or doc_name,
            "dieu": None,
            "muc": None,
            "chuong": None,
            "chunk_index": index,
            "content": chunk,
            "file_path": source.get("attachment_url") or source.get("url"),
            "url": source.get("url"),
            "attachment_url": source.get("attachment_url"),
            "source_type": "website_uneti",
            "is_active": True,
            "content_hash": content_hash,
            "updated_at": now,
            "ten_van_ban": source.get("title") or doc_name,
            "don_vi_ban_hanh": "UNETI",
            "loai_van_ban": "website",
        }
        for index, chunk in enumerate(chunks, start=1)
    ]


def _get_website_index_cache(question: str):
    key = normalize_text(question)
    cached = _WEBSITE_INDEX_CACHE.get(key)
    if not cached:
        return None

    created_at, result = cached
    if time.monotonic() - created_at > RETRIEVAL_CACHE_TTL_SECONDS:
        _WEBSITE_INDEX_CACHE.pop(key, None)
        return None

    _WEBSITE_INDEX_CACHE.move_to_end(key)
    return dict(result)


def _set_website_index_cache(question: str, result: dict):
    key = normalize_text(question)
    _WEBSITE_INDEX_CACHE[key] = (time.monotonic(), dict(result))
    _WEBSITE_INDEX_CACHE.move_to_end(key)

    while len(_WEBSITE_INDEX_CACHE) > RETRIEVAL_CACHE_MAX_ITEMS:
        _WEBSITE_INDEX_CACHE.popitem(last=False)


def index_uneti_website(question: str, debug: dict | None = None) -> dict:
    """Fetch selected UNETI web sources, chunk them, and upsert into the shared vector store."""
    cached = _get_website_index_cache(question)
    if cached is not None:
        if debug is not None:
            debug.update({
                "cache_hit": True,
                "indexed_chunks": 0,
                "indexed_sources_count": cached.get("indexed_sources_count", 0),
                "pipeline_mode": "website_to_shared_rag_cache",
            })
        return cached

    result = _search_and_extract_website_sources(question, debug)
    sources = result.get("sources", [])

    all_chunks = []
    for source in sources:
        all_chunks.extend(_build_website_chunks(source))

    indexed_chunks = index_chunks(all_chunks) if all_chunks else 0
    if indexed_chunks:
        clear_document_index_cache()

    if debug is not None:
        debug.update({
            "indexed_chunks": indexed_chunks,
            "indexed_sources_count": len(sources),
            "pipeline_mode": "website_to_shared_rag",
        })

    response = {
        "sources": sources,
        "indexed_chunks": indexed_chunks,
        "indexed_sources_count": len(sources),
    }
    _set_website_index_cache(question, response)
    return response


def _search_and_extract_website_sources(question: str, debug: dict | None = None) -> dict:
    missing_config = _validate_config()
    if missing_config:
        return _fallback_search_and_extract_website_sources(
            question,
            debug,
            reason=f"missing_config: {', '.join(missing_config)}",
        )

    try:
        from google.cloud.discoveryengine_v1beta import SearchServiceClient
    except Exception as exc:
        return _fallback_search_and_extract_website_sources(
            question,
            debug,
            reason=f"dependency_error: {exc}",
        )

    api_endpoint = (
        "discoveryengine.googleapis.com"
        if DISCOVERY_LOCATION == "global"
        else f"{DISCOVERY_LOCATION}-discoveryengine.googleapis.com"
    )
    client = SearchServiceClient(client_options={"api_endpoint": api_endpoint})
    serving_config = (
        f"projects/{DISCOVERY_PROJECT_NUMBER}/locations/{DISCOVERY_LOCATION}/"
        f"collections/{DISCOVERY_COLLECTION_ID}/engines/{DISCOVERY_ENGINE_ID}/"
        f"servingConfigs/{DISCOVERY_SERVING_CONFIG_ID}"
    )

    raw_sources = []

    try:
        response = client.search(request={
            "serving_config": serving_config,
            "query": question,
            "page_size": WEBSITE_SEARCH_TOP_K,
        })

        for result in response:
            source = _extract_document(result)
            if source:
                raw_sources.append(source)
    except Exception as exc:
        return _fallback_search_and_extract_website_sources(
            question,
            debug,
            reason=f"search_error: {exc}",
        )

    selected_sources = _dedupe_and_rerank(question, raw_sources)
    if not selected_sources:
        return _fallback_search_and_extract_website_sources(
            question,
            debug,
            reason="discovery_no_selected_sources",
        )

    answerable_sources, enriched_sources, top_source_scan_skipped = _answerable_sources_from_selected(
        selected_sources
    )

    if debug is not None:
        debug.update({
            "status": "success",
            "raw_results_count": len(raw_sources),
            "selected_sources_count": len(selected_sources),
            "answerable_sources_count": len(answerable_sources),
            "mode": "fetch_attachment_skip_scan",
            "top_source_scan_skipped": top_source_scan_skipped,
            "selected_sources": [
                {
                    "title": source.get("title"),
                    "url": source.get("url"),
                    "attachment_url": source.get("attachment_url"),
                    "score": source.get("score"),
                    "snippet_chars": len(source.get("snippet") or ""),
                    "content_chars": len(source.get("content") or ""),
                    "scan_skipped": bool(source.get("scan_skipped")),
                    "fetch_debug": source.get("fetch_debug"),
                }
                for source in enriched_sources
            ],
        })

    if not answerable_sources:
        return _fallback_search_and_extract_website_sources(
            question,
            debug,
            reason="discovery_no_answerable_sources",
        )

    return {"sources": answerable_sources}


def search_uneti_website(question: str, debug: dict | None = None) -> dict:
    answerable_sources = _search_and_extract_website_sources(question, debug).get("sources", [])
    if not answerable_sources:
        return {
            "answer": NO_WEBSITE_RESULTS_ANSWER,
            "sources": [],
        }

    prompt = _build_website_prompt(question, answerable_sources)
    answer = ask_gemini(prompt)

    sources = [
        {
            "title": source["title"],
            "url": source["url"],
            "attachment_url": source.get("attachment_url"),
            "source_type": "website_uneti",
            "score": source["score"],
            "content": source.get("content", ""),
            "preview": source.get("content", ""),
        }
        for source in answerable_sources
    ]

    return {
        "answer": answer,
        "sources": sources,
    }
