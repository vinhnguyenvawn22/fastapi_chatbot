from docx import Document

import app.controller.document_controller as documents


def test_business_docx_preserves_heading_hierarchy(monkeypatch, tmp_path):
    root = tmp_path / "document"
    business_dir = root / "nghiep_vu"
    business_dir.mkdir(parents=True)
    path = business_dir / "guide.docx"

    document = Document()
    document.add_heading("CONG TAC GIANG VIEN", level=2)
    document.add_heading("Tra cuu khoi luong", level=3)
    document.add_heading("Man Tong hop thanh toan", level=4)
    document.add_paragraph("B1: Chon nam hoc va hoc ky.")
    document.add_paragraph("B2: Xem tong so quy doi va chi tiet thanh toan.")
    document.save(path)

    monkeypatch.setattr(documents, "DOCUMENTS_DIR", str(root))
    chunks = documents.build_document_chunks("nghiep_vu/guide.docx")

    assert len(chunks) == 1
    assert chunks[0]["document_type"] == "business_document"
    assert chunks[0]["heading"] == "Man Tong hop thanh toan"
    assert chunks[0]["section_path"] == (
        "CONG TAC GIANG VIEN > Tra cuu khoi luong > Man Tong hop thanh toan"
    )
    assert "B2: Xem tong so quy doi" in chunks[0]["content"]


def test_business_docx_groups_consecutive_workflow_steps(monkeypatch, tmp_path):
    root = tmp_path / "document"
    business_dir = root / "nghiep_vu"
    business_dir.mkdir(parents=True)
    path = business_dir / "survey.docx"

    document = Document()
    document.add_heading("Khao sat noi bo", level=3)
    document.add_heading("Buoc 1: Dang nhap", level=4)
    document.add_paragraph("Dang nhap bang tai khoan ca nhan.")
    document.add_heading("Buoc 2: Mo phieu", level=4)
    document.add_paragraph("Chon phieu khao sat can tra loi.")
    document.add_heading("Buoc 3: Nop phieu", level=4)
    document.add_paragraph("Nhan nut Nop khao sat.")
    document.save(path)

    monkeypatch.setattr(documents, "DOCUMENTS_DIR", str(root))
    chunks = documents.build_document_chunks("nghiep_vu/survey.docx")

    assert len(chunks) == 1
    assert chunks[0]["section_type"] == "business_workflow"
    assert chunks[0]["section_path"] == "Khao sat noi bo"
    assert all(f"Buoc {step}" in chunks[0]["content"] for step in (1, 2, 3))


def test_regulation_keeps_article_chunking(monkeypatch, tmp_path):
    root = tmp_path / "document"
    root.mkdir()
    path = root / "Quy dinh dao tao.docx"

    document = Document()
    document.add_paragraph("Dieu 1. Pham vi ap dung")
    document.add_paragraph("Noi dung dieu mot.")
    document.add_paragraph("Dieu 2. Dieu kien")
    document.add_paragraph("Noi dung dieu hai.")
    document.save(path)

    monkeypatch.setattr(documents, "DOCUMENTS_DIR", str(root))
    chunks = documents.build_document_chunks("Quy dinh dao tao.docx")

    assert [chunk["dieu"] for chunk in chunks] == [1, 2]
    assert all(chunk["document_type"] == "regulation" for chunk in chunks)
